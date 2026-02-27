"""
Add new sections to the WOLF Engine Overview document.
Run AFTER generate_overview.py to patch the docx with new modules.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'WOLF_Engine_Complete_Overview.docx')
doc = Document(path)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

# ============================================================
# Insert new sections at the end (before saving)
# We'll add them as appendix-style sections
# ============================================================

doc.add_page_break()

# ── FCFF THREE-WAY VERIFICATION ──
doc.add_heading('Appendix A: FCFF Three-Way Cross-Verification', level=1)
doc.add_paragraph('File: src/utils/valuationEngine.ts - calculateFCFFVerification()')
doc.add_paragraph(
    'All three methods must produce identical FCFF results within a 1% tolerance. '
    'This is a critical audit check that ensures internal consistency of financial data.'
)
add_table(
    ['Method', 'Formula', 'Route'],
    [
        ['Method 1 (NOPAT)', 'EBIT x (1-t) + D&A - CapEx - Delta WC', 'Primary'],
        ['Method 2 (EBITDA)', 'EBITDA x (1-t) + D&A x t - CapEx - Delta WC', 'Cross-check'],
        ['Method 3 (Net Income)', 'NI + Interest x (1-t) + D&A - CapEx - Delta WC', 'Cross-check'],
    ]
)
doc.add_paragraph(
    'The FCFFVerification interface returns: method1, method2, method3 values, '
    'allMatch boolean, and tolerance percentage. The UI displays a reconciliation '
    'table in the FCFFReconciliation.tsx component with color-coded match indicators.'
)

# ── DDM (DIVIDEND DISCOUNT MODELS) ──
doc.add_heading('Appendix B: Dividend Discount Models (DDM)', level=1)
doc.add_paragraph('File: src/utils/valuationEngine.ts - calculateDDM()')
doc.add_paragraph(
    'Three DDM variants are implemented. All discount at Cost of Equity (Ke), NOT WACC, '
    'since DDM values equity directly. DDM is only applicable when dividendsPerShare > 0.'
)
add_table(
    ['Model', 'Formula', 'Use Case'],
    [
        ['Gordon Growth', 'DPS x (1+g) / (Ke - g)', 'Stable dividend growers'],
        ['Two-Stage', 'PV(high-growth dividends) + PV(stable terminal)', 'Companies transitioning growth rates'],
        ['H-Model', 'DPS x (1+gL) / (Ke-gL) + DPS x H x (gS-gL) / (Ke-gL)', 'Gradually declining growth'],
    ]
)
doc.add_paragraph('DDMResult interface: gordonGrowth, twoStage, hModel (each number|null), applicable boolean, message string.')

# ── EAS COMPLIANCE ──
doc.add_heading('Appendix C: EAS / IFRS Compliance Module', level=1)
doc.add_paragraph('File: src/utils/easModules.ts')
doc.add_paragraph(
    'Four Egyptian Accounting Standards compliance modules with IFRS equivalents. '
    'Each function is pure, taking inputs and returning structured results.'
)

doc.add_heading('C.1 EAS 48 (IFRS 16) - Lease Adjustments', level=2)
doc.add_paragraph('Function: calculateLeaseAdjustment(inputs, currentEBITDA, currentTotalDebt)')
items = [
    'Capitalizes operating leases onto the balance sheet as Right-of-Use (ROU) assets',
    'ROU Asset = PV of remaining lease payments at Incremental Borrowing Rate (IBR)',
    'Formula: PV = PMT x [(1 - (1+r)^-n) / r]',
    'EBITDA add-back = full annual lease payment (reclassified from operating to financing)',
    'Debt increase = lease liability (equal to ROU asset value at inception)',
    'Depreciation = ROU / remaining term (straight-line)',
    'Interest expense = lease liability x IBR (Year 1)',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('C.2 EAS 31 (IAS 1) - Normalized Earnings', level=2)
doc.add_paragraph('Function: calculateNormalizedEarnings(items, reportedNI, taxRate, shares)')
items = [
    'Adds back / deducts non-recurring items from reported Net Income',
    'Tax-affected items are adjusted at the marginal tax rate',
    'Common add-backs: impairment charges, restructuring costs, legal settlements',
    'Common deductions: gain on asset sale, insurance proceeds, one-time subsidies',
    'Returns: grossAdjustment, taxEffect, netAdjustment, normalizedNI, normalizedEPS',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('C.3 EAS 12 (IAS 12) - Deferred Tax in EV Bridge', level=2)
doc.add_paragraph('Function: calculateDeferredTaxAdjustment(inputs, baseEquityValue)')
items = [
    'Net DTA (DTA > DTL) adds to equity value (future tax savings)',
    'Net DTL (DTL > DTA) reduces equity value (future tax obligations)',
    'Adjusted Equity = Base Equity Value + Net DTA',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('C.4 EAS 23 (IAS 33) - Basic & Diluted EPS', level=2)
doc.add_paragraph('Function: calculateEPS(netIncome, basicShares, dilutedShares, convertibleInterest)')
items = [
    'Basic EPS = Net Income / Weighted Average Basic Shares',
    'Diluted EPS = (NI + Convertible Interest) / (Basic + Dilutive Shares)',
    'Anti-dilution check: if Diluted EPS > Basic EPS, dilutive securities are excluded',
    'Returns: basicEPS, dilutedEPS, dilutionPercent, isAntiDilutive flag',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

# ── CONFIDENCE SCORING ──
doc.add_heading('Appendix D: Model Confidence Score (0-100)', level=1)
doc.add_paragraph('File: src/utils/confidenceScore.ts - calculateConfidenceScore()')
doc.add_paragraph(
    'Three-pillar scoring system that evaluates the reliability of the valuation model. '
    'Total score maps to letter grade: A (85+), B (70+), C (55+), D (40+), F (<40).'
)
add_table(
    ['Pillar', 'Max Points', 'Checks'],
    [
        ['Data Quality', '30 pts', 'D&A>0 (5), CapEx>0 (5), Interest consistency (5), EBITDA>0 (5), Revenue>0 (5), Shares>0 (5)'],
        ['Assumption Reasonableness', '40 pts', 'Rev growth<=25% (8), EBITDA margin<=50% (8), Terminal g<GDP (8), WACC 15-35% (8), FCFF reconciles (8)'],
        ['Model Robustness', '30 pts', 'TV<80% of EV (10), Value within +/-50% of market (10), DCF/DDM agreement (5), Interest coverage>1.5x (5)'],
    ]
)
doc.add_paragraph(
    'Returns: totalScore, grade (A-F), per-pillar scores, detailed breakdown array '
    'with per-check scores, and a narrative explanation of concerns.'
)

# ── INPUT VALIDATION ──
doc.add_heading('Appendix E: Input Validation Engine', level=1)
doc.add_paragraph('File: src/utils/inputValidation.ts - validateInputs()')
doc.add_paragraph('Three severity levels with distinct behavior:')

doc.add_heading('Hard Blocks (Must Fix)', level=2)
blocks = [
    'Revenue <= 0: Cannot perform valuation without revenue',
    'Shares Outstanding <= 0: Cannot calculate per-share values',
    'WACC <= Terminal Growth: Gordon Growth Model produces infinite/negative TV',
    'WACC <= 0: Invalid discount rate',
    'Projection Years < 1 or > 20: Out of reasonable range',
    'Stock Price <= 0: Cannot calculate upside/downside',
    'EBITDA Margin <= 0% or > 95%: Unrealistic operating profile',
    'Revenue Growth > 100%: Unreasonably aggressive',
    'Tax Rate < 0% or > 60%: Out of legal range',
]
for item in blocks:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Soft Warnings (Review Recommended)', level=2)
warnings = [
    'Terminal Value > 80% of EV: Model relies too heavily on terminal assumptions',
    'Implied price deviates > 50% from market: May indicate aggressive assumptions',
    'Negative FCF in base year: DCF may extrapolate losses',
    'Interest Coverage < 1.5x: Company may face debt service issues',
    'D&A = 0 with positive PP&E: May understate reinvestment needs',
    'CapEx = 0: FCFF may be overstated',
    'Revenue growth > 25%: Aggressive for most companies',
    'Debt/Equity > 3.0: Highly leveraged',
]
for item in warnings:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Edge Case Auto-Adjustments', level=2)
edge = [
    'Zero interest expense with positive debt: Impute at ~2.8% of total debt',
    'COGS = 0 for banks: Recalculate as Revenue - Interest Expense',
    'Negative equity: Flag but allow calculation to proceed',
    'Missing D&A: Default to 0 with warning',
]
for item in edge:
    doc.add_paragraph(item, style='List Bullet')

# ── JSON EXPORT ──
doc.add_heading('Appendix F: JSON Export (CFA-Grade)', level=1)
doc.add_paragraph('File: src/utils/jsonExport.ts')
doc.add_paragraph(
    'Generates a fully-compliant JSON export with complete metadata and audit trail. '
    'Engine version: 4.0.0. Schema follows Section 8.4 specification.'
)
doc.add_paragraph('The JSON structure contains three top-level sections:')
add_table(
    ['Section', 'Contents'],
    [
        ['metadata', 'engine_version, generation_date, currency, company_name, capm_method, rf_instrument, rf_date, discounting_convention, accounting_standard'],
        ['inputs', 'All financial data fields, all assumption parameters (25+ fields)'],
        ['calculated', 'WACC result, DCF result with full projections and EV bridge, DDM result, comparable multiples, sensitivity matrix, financial ratios, FCFF verification, scenario analysis'],
    ]
)

# ── UPDATED COMPONENT LIST ──
doc.add_heading('Appendix G: Updated Valuation Section Components', level=1)
doc.add_paragraph('File: src/components/valuation/sections/ (14 components)')
add_table(
    ['Component', 'Purpose'],
    [
        ['ValidationAlerts.tsx', 'Market-specific validation warnings with severity levels'],
        ['ValuationStyleSelector.tsx', 'Conservative/Moderate/Aggressive with side-by-side DCF comparison'],
        ['MarketVsFundamental.tsx', 'Market price vs intrinsic value with key statistics'],
        ['ScenarioAnalysis.tsx', 'Bear/Base/Bull with sensitivity matrix and Monte Carlo'],
        ['ValuationSummaryCards.tsx', 'DCF, DDM, Comps, and Blended value cards'],
        ['BlendedWeightSlider.tsx', 'Adjustable DCF/Comps weighting with live preview'],
        ['ComparableBreakdown.tsx', 'Individual method implied prices (P/E, EV/EBITDA, P/S, P/B)'],
        ['BaseYearFCF.tsx', 'Base year free cash flow breakdown'],
        ['DCFProjectionsTable.tsx', 'Year-by-year revenue, EBITDA, D&A, EBIT, NOPAT, CapEx, WC, FCFF'],
        ['KeyMetricsGrid.tsx', '30+ financial ratios in card format'],
        ['ReverseDCFSection.tsx', 'Market-implied growth rate analysis'],
        ['QualityScorecard.tsx', '4-category quality assessment with letter grade'],
        ['FCFFReconciliation.tsx', 'Three-method FCFF cross-verification display'],
        ['EASComplianceSection.tsx', 'EAS 48/31/12/23 compliance panels with calculations'],
    ]
)

# ── Save ──
doc.save(path)
print(f'Updated document saved to: {path}')
