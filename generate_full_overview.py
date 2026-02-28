"""
WOLF Valuation Engine — Complete Overview Word Document Generator
Generates a unified, comprehensive document covering ALL engine modules.
All values verified against actual source code as of Feb 2026.
Run: python generate_full_overview.py
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

from doc_sections import build_title_page, build_toc, build_executive_summary, build_tech_stack

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
for i in range(1, 5):
    doc.styles[f'Heading {i}'].font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

def bullets(items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# ═══════════════════════════════
# BUILD SECTIONS 1-2
# ═══════════════════════════════
build_title_page(doc, Pt, RGBColor, WD_ALIGN_PARAGRAPH)
build_toc(doc, Pt)
build_executive_summary(doc, add_table)
build_tech_stack(doc, add_table)

# ═══════════════════════════════
# 3. PROJECT STRUCTURE
# ═══════════════════════════════
doc.add_heading('3. Project Structure & File Map', level=1)
doc.add_paragraph('Complete directory tree with all source files:')
structure = """ValuationEngine/
+-- index.html                    # SPA entry point
+-- package.json                  # Dependencies & scripts
+-- vite.config.ts                # Vite build config (React + Tailwind + singlefile)
+-- tsconfig.json                 # TypeScript config (strict, ES2020, React JSX)
+-- .env.example                  # VITE_FMP_API_KEY template
+-- src/
    +-- main.tsx                  # React DOM bootstrap with ErrorBoundary
    +-- App.tsx                   # Main orchestrator (176 lines)
    +-- index.css                 # Root CSS imports
    +-- vite-env.d.ts             # Vite type declarations
    |
    +-- types/
    |   +-- financial.ts          # ALL TypeScript interfaces (465 lines, 55+ types)
    |
    +-- constants/
    |   +-- initialData.ts        # Default FinancialData & ValuationAssumptions
    |   +-- marketDefaults.ts     # USA/Egypt params, tax categories, industry multiples (96 lines)
    |   +-- valuationStyles.ts    # Conservative/Moderate/Aggressive presets (44 lines)
    |
    +-- hooks/
    |   +-- useFinancialData.ts   # Central state: financialData + assumptions + comparables + history
    |   +-- useValuationCalculations.ts  # Memoized: WACC, FCFF, DCF, DDM, Comps, Ratios, Sensitivity, etc.
    |   +-- useHistory.ts         # Undo/Redo stack management
    |   +-- useTheme.ts           # Dark/Light mode toggle
    |   +-- useKeyboardShortcuts.ts  # Ctrl+Z, Ctrl+Y, Ctrl+S
    |
    +-- services/
    |   +-- stockAPI.ts           # FMP Stable API: profile, IS, BS, CF, ratios + peer fetching
    |   +-- yahooFinanceAPI.ts    # Yahoo Finance v10 via CORS proxy chain (Egyptian stocks)
    |
    +-- utils/
    |   +-- valuationEngine.ts    # CORE ENGINE (1,098 lines): WACC, FCFF, DCF, DDM, Comps, Blended,
    |   |                         #   Ratios, Sensitivity, Scenarios, ReverseDCF, AutoCalc, Validation
    |   +-- advancedAnalysis.ts   # Monte Carlo, Sector Benchmarking, Quality Scorecard (28KB)
    |   +-- easModules.ts         # EAS/IFRS compliance: Lease, NormalizedEarnings, DeferredTax, EPS (232 lines)
    |   +-- confidenceScore.ts    # 3-pillar confidence scoring 0-100 (186 lines)
    |   +-- inputValidation.ts    # Hard blocks / Warnings / Edge cases (353 lines)
    |   +-- jsonExport.ts         # CFA-grade JSON export with audit trail (144 lines)
    |   +-- pdfExport.ts          # PDF research report (33KB)
    |   +-- excelExport.ts        # Multi-sheet Excel export (45KB)
    |   +-- excelExportPro.ts     # Excel with live formulas (38KB)
    |   +-- formatters.ts         # formatCurrency, formatPercent, formatMultiple, etc. (6KB)
    |   +-- industryMapping.ts    # Ticker-to-sector mapping
    |   +-- valuation.ts          # Legacy calc helpers (retained for compat)
    |   +-- cn.ts                 # clsx + tailwind-merge utility
    |   +-- calculations/
    |       +-- dcf.ts            # DCF projection & valuation functions
    |       +-- comparables.ts    # Comparable company functions
    |       +-- metrics.ts        # Key financial metrics & recommendations
    |       +-- scenarios.ts      # Bear/Base/Bull scenario functions
    |
    +-- workers/
    |   +-- monteCarlo.worker.ts  # Dedicated thread for Monte Carlo
    |
    +-- components/               # 40+ components
        +-- layout/               # Header, CompanyHeader, TabNavigation, Footer
        +-- shared/               # ErrorBoundary, LoadingFallback, InputField, CalculationAuditTrail
        +-- input/InputTab.tsx    # All data entry forms
        +-- charts/ChartsTab.tsx  # Revenue projection, football field charts
        +-- valuation/
        |   +-- ValuationTab.tsx  # Tab orchestrator
        |   +-- sections/ (15 files)
        +-- StockSearch, FinancialInputForm, AssumptionsPanel, ComparableCompanies,
            DCFResults, KeyMetrics, ValuationSummary, FootballFieldChart, ScenarioToggle,
            DebouncedInput, Tooltip, AIReport, QuickStartGuide, HowToBuildModal,
            APIKeyModal, UserAuth, WolfLogo"""
p = doc.add_paragraph()
r = p.add_run(structure)
r.font.name = 'Consolas'
r.font.size = Pt(7)
doc.add_page_break()

# ═══════════════════════════════
# 4. DATA MODELS
# ═══════════════════════════════
doc.add_heading('4. Data Models & TypeScript Interfaces', level=1)
doc.add_paragraph('File: src/types/financial.ts (465 lines, 55+ type definitions). Every interface documented.')

doc.add_heading('4.1 FinancialData (Central Data Object)', level=2)
add_table(['Field', 'Type', 'Description'], [
    ['companyName', 'string', 'Company display name'],
    ['ticker', 'string', 'Stock ticker (e.g., AAPL, COMI.CA)'],
    ['sharesOutstanding', 'number', 'Total shares outstanding'],
    ['currentStockPrice', 'number', 'Current market price per share'],
    ['dividendsPerShare', 'number', 'Annual DPS (used by DDM models)'],
    ['lastReportedDate?', 'string', 'Last fiscal year end date from API'],
    ['sector?', 'string', 'Company sector from API profile'],
    ['fiscalYearEnd?', 'string', 'Fiscal year end (dec/jun/mar/sep)'],
    ['incomeStatement', 'IncomeStatement', '10 fields: revenue through amortization'],
    ['balanceSheet', 'BalanceSheet', '25 fields including optional EAS fields'],
    ['cashFlowStatement', 'CashFlowStatement', '6 fields including otherFinancingActivities'],
    ['historicalData?', 'HistoricalYear[]', 'Multi-period data (2-5 prior years)'],
])

doc.add_heading('4.2 BalanceSheet (25 Fields)', level=2)
doc.add_paragraph('Includes new optional fields for enhanced EV bridge and analysis:')
add_table(['Category', 'Fields'], [
    ['Current Assets', 'cash, marketableSecurities, accountsReceivable, inventory, otherCurrentAssets, totalCurrentAssets'],
    ['Non-Current Assets', 'propertyPlantEquipment, longTermInvestments, goodwill, intangibleAssets, otherNonCurrentAssets, totalAssets'],
    ['Current Liabilities', 'accountsPayable, shortTermDebt, otherCurrentLiabilities, totalCurrentLiabilities'],
    ['Non-Current Liabilities', 'longTermDebt, otherNonCurrentLiabilities, totalLiabilities'],
    ['Equity', 'totalEquity, retainedEarnings? (for Altman Z), minorityInterest? (EV bridge), preferredEquity? (EV bridge)'],
    ['EAS-Specific', 'endOfServiceProvision? (EAS 42 — treated as debt-like for EV bridge)'],
])

doc.add_heading('4.3 ValuationAssumptions (30+ Fields)', level=2)
add_table(['Category', 'Fields'], [
    ['Core DCF', 'discountRate (WACC), terminalGrowthRate, projectionYears (3/5/7/10), revenueGrowthRate, taxRate'],
    ['CAPM', 'riskFreeRate, marketRiskPremium, beta, capmMethod (A/B), betaType (raw/adjusted/relevered), taxCategory'],
    ['Method B', 'rfUS (US Treasury), countryRiskPremium, egyptInflation, usInflation'],
    ['Cost of Debt', 'costOfDebt (pre-tax %)'],
    ['Projection Drivers', 'ebitdaMargin, daPercent, capexPercent, deltaWCPercent, useConstantDrivers (bool)'],
    ['Terminal Value', 'terminalMethod (gordon_growth / exit_multiple), exitMultiple'],
    ['Discounting', 'discountingConvention (end_of_year / mid_year)'],
    ['DDM', 'dps, ddmHighGrowth, ddmStableGrowth, ddmHighGrowthYears'],
    ['Scenarios', 'bearProbability, baseProbability, bullProbability'],
    ['Metadata', 'rfDate (date stamp of risk-free rate), marginImprovement (legacy)'],
])

doc.add_heading('4.4 Key Result Interfaces', level=2)
add_table(['Interface', 'Key Fields', 'Lines'], [
    ['WACCResult', 'costOfEquity, afterTaxCostOfDebt, equityWeight, debtWeight, wacc, marketCap, totalDebt, capmMethod, keUSD?', '199-211'],
    ['FCFFVerification', 'method1, method2, method3, allMatch, tolerance (0.001)', '217-223'],
    ['DCFProjection', 'year, revenue, ebitda, dAndA, ebit, nopat, capex, deltaWC, freeCashFlow, discountFactor, presentValue', '181-193'],
    ['DCFResult', 'sumOfPVs, terminalValue, pvOfTerminal, tvPercent, EV, netDebt, equityValue, impliedSharePrice, upside, marginOfSafety, verdict, discountingConvention', '229-242'],
    ['DDMResult', 'gordonGrowth, twoStage, hModel (each number|null), applicable, message?', '248-254'],
    ['SensitivityMatrix', 'waccAxis[], growthAxis[], cells[][] (SensitivityCell: wacc, growth, impliedPrice, upside, color, isBaseCase)', '269-273'],
    ['ScenarioAnalysis', 'bear, base, bull (ScenarioCase: revGrowth, ebitdaMargin, termGrowth, wacc, probability, intrinsicValue), weightedValue', '288-293'],
    ['ReverseDCFResult', 'impliedGrowthRate, baseGrowthRate, growthGap, marketExpectation (aggressive/reasonable/conservative), narrative', '299-305'],
    ['FinancialRatios', '30+ fields: margins (4), returns (3), leverage (3), liquidity (2), efficiency (4), value creation (2), quality (2), multiples (11)', '322-368'],
    ['ValidationAlert', 'type (error/warning/info), message, field?, blocking?', '311-316'],
    ['ValuationJSON', 'metadata (9 fields) + inputs (25+ fields) + calculated (wacc + dcf with projections + ev_bridge + ddm + multiples + sensitivity + ratios + fcff_verification + scenarios)', '374-423'],
    ['HistoricalYear', 'year, revenue, netIncome, totalAssets, totalEquity, operatingCashFlow, capex, grossMargin, longTermDebt, currentRatio, sharesOutstanding', '65-77'],
])

doc.add_heading('4.5 Union Types & Enums', level=2)
add_table(['Type', 'Values', 'Purpose'], [
    ['MarketRegion', '"USA" | "Egypt"', 'Dual-market selection'],
    ['CAPMMethod', '"A" | "B"', 'Method A = Local CAPM, B = USD Build-Up'],
    ['BetaType', '"raw" | "adjusted" | "relevered"', 'Beta calculation method'],
    ['DiscountingConvention', '"end_of_year" | "mid_year"', 'Cash flow timing'],
    ['TerminalValueMethod', '"gordon_growth" | "exit_multiple"', 'Terminal value approach'],
    ['EgyptTaxCategory', '"standard" | "oil_gas" | "suez_canal" | "free_zone" | "custom"', 'Egyptian tax brackets'],
])
doc.add_page_break()

# ═══════════════════════════════
# 5. CORE VALUATION ENGINE
# ═══════════════════════════════
doc.add_heading('5. Core Valuation Engine', level=1)
doc.add_paragraph(
    'File: src/utils/valuationEngine.ts (1,098 lines). Every formula verified against CFA specification. '
    'All functions are pure — no React dependencies.'
)

doc.add_heading('5.1 WACC Calculation (Dual CAPM)', level=2)
doc.add_paragraph('Functions: calculateCostOfEquity(inputs) + calculateWACC(inputs) -> WACCResult')
doc.add_paragraph('Formula: WACC = (E/V) * Ke + (D/V) * Kd * (1 - t)')
doc.add_paragraph('Capital structure uses Market Cap (NOT book equity) for E/V weights.')

doc.add_heading('Method A — Local Currency CAPM (Default for EGP)', level=3)
bullets([
    'Ke = Rf(Egypt) + Beta * Mature_Market_ERP',
    'NO Country Risk Premium added — Egyptian Rf already embeds country risk',
    'Default: Rf = 22.0% (10Y Egyptian Govt Bond), ERP = 5.5% (mature market only)',
])
doc.add_heading('Method B — USD Build-Up', level=3)
bullets([
    'Ke(USD) = Rf(US) + Beta * Mature_ERP + CRP',
    'Ke(EGP) = (1 + Ke_USD) * (1 + Egypt_Inflation) / (1 + US_Inflation) - 1',
    'Fisher equation converts USD cost of equity to EGP terms',
    'Default: Rf(US)=4.5%, ERP=5.5%, CRP=7.5% (Damodaran), Egypt Infl=28%, US Infl=3%',
])
doc.add_heading('Beta Adjustment (Section 3.3)', level=3)
bullets([
    'Raw: Use as-is from API or user input',
    'Adjusted (Bloomberg): Adjusted Beta = (2/3 * Raw Beta) + (1/3 * 1.0)',
    'Relevered: BetaL = BetaU * [1 + (1-t) * (D/E)] — handled externally',
])

doc.add_heading('5.2 FCFF Three-Way Cross-Verification (Section 4)', level=2)
doc.add_paragraph('Function: calculateFCFFVerification() — All 3 must match within 0.001 tolerance')
add_table(['Method', 'Full Formula'], [
    ['1: NOPAT Route', 'FCFF = EBIT * (1-t) + D&A - CapEx - DeltaWC'],
    ['2: EBITDA Route', 'FCFF = EBITDA * (1-t) + D&A * t - CapEx - DeltaWC'],
    ['3: Net Income Route', 'FCFF = Net Income + Interest * (1-t) + D&A - CapEx - DeltaWC'],
])
doc.add_paragraph('WARNING: FCFF must NEVER subtract Interest Expense directly. DeltaWC increase = cash OUTFLOW (subtracted).')

doc.add_heading('5.3 DCF Projections (Section 5.1)', level=2)
doc.add_paragraph('Function: calculateDCFProjections() — Year-by-year with proper FCFF formula')
doc.add_paragraph('Projection chain per year:')
add_table(['Step', 'Formula'], [
    ['Revenue(t)', 'Revenue(t-1) * (1 + Revenue_Growth)'],
    ['EBITDA(t)', 'Revenue(t) * EBITDA_Margin'],
    ['D&A(t)', 'Revenue(t) * DA_Percent'],
    ['EBIT(t)', 'EBITDA(t) - D&A(t)'],
    ['NOPAT(t)', 'EBIT(t) * (1 - Tax_Rate)'],
    ['CapEx(t)', 'Revenue(t) * CapEx_Percent'],
    ['DeltaWC(t)', '(Revenue(t) - Revenue(t-1)) * DeltaWC_Percent'],
    ['FCFF(t)', 'NOPAT + D&A - CapEx - DeltaWC'],
    ['Discount Factor', '(1 + WACC)^period  where period = year-0.5 (mid-year) or year (end-of-year)'],
    ['Present Value', 'FCFF / Discount Factor'],
])

doc.add_heading('5.4 DCF Valuation (Section 5.2-5.4)', level=2)
doc.add_paragraph('Function: calculateDCFValue() — Terminal Value + EV-to-Equity Bridge')
bullets([
    'Terminal Value (Gordon Growth): TV = FCFF_N * (1+g) / (WACC - g)',
    'Terminal Value (Exit Multiple): TV = EBITDA_N * Exit Multiple',
    'Hard block: If WACC <= terminal growth, TV = 0 (mathematically invalid)',
    'Enterprise Value = Sum of PV(FCFFs) + PV(Terminal Value)',
    'Equity Value = EV - Total Debt + Cash (NO MAX(0) floor — can be negative)',
    'Implied Share Price = Equity Value / Shares Outstanding',
    'Upside = (Implied - Current) / Current * 100%',
    'Margin of Safety = (Implied - Current) / Implied * 100%',
    'Verdict: UNDERVALUED (>10% upside), OVERVALUED (>10% downside), FAIRLY VALUED (within +/-10%)',
    'Terminal Value as % of EV tracked for confidence scoring',
])

doc.add_heading('5.5 Dividend Discount Models (Section 6)', level=2)
doc.add_paragraph('Function: calculateDDM() — Discounts at Ke (cost of equity), NOT WACC')
doc.add_paragraph('Applicability: Only when dividendsPerShare > 0 AND net income > 0. DPS calculated from actual dividendsPaid / sharesOutstanding.')
add_table(['Model', 'Formula', 'Use Case'], [
    ['6.1 Gordon Growth', 'P = DPS * (1+g) / (Ke - g)', 'Stable dividend growers, single perpetual growth rate'],
    ['6.2 Two-Stage', 'P = Sum[DPS*(1+gH)^t / (1+Ke)^t] for t=1..n + PV(Terminal)', 'Companies transitioning from high to stable growth'],
    ['6.3 H-Model', 'P = DPS*(1+gL)/(Ke-gL) + DPS*H*(gH-gL)/(Ke-gL), H=n/2', 'Gradually declining growth (linear decay)'],
])

doc.add_heading('5.6 Comparable Company Valuation', level=2)
doc.add_paragraph('Function: calculateComparableValue() — Median peer multiples')
add_table(['Method', 'Formula'], [
    ['P/E', 'EPS * Median Peer P/E'],
    ['EV/EBITDA', '(EBITDA * Median Peer EV/EBITDA - Total Debt + Cash) / Shares'],
    ['P/S', 'Revenue per Share * Median Peer P/S'],
    ['P/B', 'Book Value per Share * Median Peer P/B'],
])
doc.add_paragraph('Average = mean of valid (>0) implied prices. No MAX(0) floor on equity calculations (EV/EBITDA can produce negative price for over-leveraged companies).')

doc.add_heading('5.7 Blended Valuation & Recommendations', level=2)
doc.add_paragraph('Function: calculateBlendedValuation() — Default 60% DCF / 40% Comps. Falls back to 100% DCF if no valid comparables.')
add_table(['Upside Range', 'Verdict'], [
    ['> +10%', 'UNDERVALUED (green)'],
    ['-10% to +10%', 'FAIRLY VALUED (yellow)'],
    ['< -10%', 'OVERVALUED (red)'],
])

doc.add_heading('5.8 Financial Ratios (30+ — Section 7)', level=2)
doc.add_paragraph('Function: calculateFinancialRatios(financialData, wacc)')
add_table(['Category', 'Ratios'], [
    ['Profitability (4)', 'Gross Margin, EBITDA Margin, EBIT Margin, Net Margin'],
    ['Returns (3)', 'ROE, ROA, ROIC (NOPAT / Invested Capital)'],
    ['Leverage (3)', 'Debt/Equity, Net Debt/EBITDA, Interest Coverage'],
    ['Liquidity (2)', 'Current Ratio, Quick Ratio (cash+AR/CL)'],
    ['Efficiency (4)', 'DSO, DIO, DPO, Cash Conversion Cycle (DSO+DIO-DPO)'],
    ['Value Creation (2)', 'EVA (NOPAT - IC*WACC), ROIC-WACC Spread'],
    ['Quality (2)', 'Altman Z-Score (5-factor: WC/TA + RE/TA + EBIT/TA + MCap/TL + Rev/TA), Piotroski F-Score (TODO)'],
    ['Multiples (11)', 'Trailing P/E, Forward P/E, EV/EBITDA, EV/EBIT, EV/Revenue, P/B, P/S, P/CF, PEG, Dividend Yield, Earnings Yield, FCF Yield'],
])

doc.add_heading('5.9 Sensitivity Analysis (Section 8)', level=2)
doc.add_paragraph('Function: generateSensitivityMatrix() — Dynamic 5x5 WACC vs Terminal Growth')
bullets([
    'WACC axis: base +/- 2%, +/- 4% (centered on calculated WACC)',
    'Growth axis: base +/- 1.5%, +/- 3% (floored at 4% for Egypt, capped below WACC-1%)',
    'Each cell recalculates full DCF with that WACC/growth combination',
    'Color-coded: green (>10% upside), yellow (-10% to +10%), red (<-10%)',
    'Invalid combinations (WACC <= growth) show $0 / -100% / red',
    'Base case cell highlighted (isBaseCase flag)',
])

doc.add_heading('5.10 Scenario Analysis (Section 8.2)', level=2)
doc.add_paragraph('Function: calculateScenarioAnalysis() — Probability-weighted')
add_table(['Scenario', 'Rev Growth Delta', 'Margin Delta', 'Terminal g Delta', 'WACC Delta', 'Default Prob'], [
    ['Bear', '-4%', '-2%', '-1.5%', '+3%', '25%'],
    ['Base', '0', '0', '0', '0', '50%'],
    ['Bull', '+6%', '+3%', '+1.5%', '-2.5%', '25%'],
])
doc.add_paragraph('Weighted Value = Sum(scenario intrinsic value * probability). Probabilities user-adjustable (must sum to 100%).')

doc.add_heading('5.11 Reverse DCF (Section 8.3)', level=2)
doc.add_paragraph('Function: calculateReverseDCF() — Binary search (100 iterations)')
doc.add_paragraph('Finds revenue growth rate that produces current stock price as DCF output. Categorizes market expectation as "aggressive" (implied > base + 3%), "reasonable" (within 3%), or "conservative" (implied < base - 3%).')

doc.add_heading('5.12 Auto-Calculate Assumptions', level=2)
doc.add_paragraph('Function: calculateAssumptionsFromData(financialData, marketRegion, apiBeta?)')
doc.add_paragraph('Derives: beta, taxRate, revenueGrowth, fcfMargin, costOfDebt, calculatedWACC from financial statements. Called when stock data is fetched from API.')

doc.add_heading('5.13 Input Validation (Inline)', level=2)
doc.add_paragraph('Function: validateInputs(financialData, assumptions, marketRegion) -> ValidationAlert[]')
doc.add_paragraph('20+ validation checks returning error/warning/info alerts. Separate from inputValidation.ts module.')
doc.add_page_break()

# ═══════════════════════════════
# 6. ADVANCED ANALYSIS
# ═══════════════════════════════
doc.add_heading('6. Advanced Analysis Module', level=1)
doc.add_paragraph('File: src/utils/advancedAnalysis.ts (28,834 bytes)')

doc.add_heading('6.1 Monte Carlo Simulation', level=2)
doc.add_paragraph('5,000 DCF simulations with Gaussian-distributed parameters:')
add_table(['Parameter', 'Base', 'Std Dev'], [
    ['Revenue Growth', 'User assumption', '+/- 4%'],
    ['WACC', 'User assumption', '+/- 1.5%'],
    ['Terminal Growth', 'User assumption', '+/- 0.8%'],
    ['Margin Improvement', 'User assumption', '+/- 0.5%'],
])
doc.add_paragraph('Outputs: mean, median, std dev, percentiles (5th/25th/75th/95th), probability of exceeding current and base case prices, 20-bucket histogram. Also available as Web Worker (src/workers/monteCarlo.worker.ts).')

doc.add_heading('6.2 Sector Benchmarking', level=2)
doc.add_paragraph('Compares vs sector averages (TECH, FINANCE, DEFAULT) across 10 metrics: Gross Margin, Operating Margin, Net Margin, ROE, ROA, FCF Margin, Debt/Equity, Current Ratio, P/E, EV/EBITDA. Each scored top/above/average/below/bottom. Overall score 0-100 with rating and insight narrative.')

doc.add_heading('6.3 Quality Scorecard', level=2)
add_table(['Category (out of 10)', 'Factors'], [
    ['Economic Moat', 'Gross Margin, Operating Margin, ROE, Market Position (revenue size)'],
    ['Financial Health', 'Current Ratio, Debt/Equity, Interest Coverage, Cash/Debt ratio'],
    ['Growth & Profitability', 'Net Margin, FCF Generation, Earnings Quality (OCF > NI)'],
    ['Capital Allocation', 'CapEx/Revenue, Dividend Payout, FCF Conversion, Cash Accumulation'],
])
doc.add_paragraph('Letter grade A+ to F with quality premium/discount adjustment: +15% (A+), +10% (A), +5% (B+), 0% (B), -5% (C), -10% (D), -15% (F).')
doc.add_page_break()

# ═══════════════════════════════
# 7. EAS COMPLIANCE
# ═══════════════════════════════
doc.add_heading('7. EAS / IFRS Compliance Module', level=1)
doc.add_paragraph('File: src/utils/easModules.ts (232 lines). Four Egyptian Accounting Standards with IFRS equivalents. All pure functions.')

doc.add_heading('7.1 EAS 48 (IFRS 16) — Lease Adjustments', level=2)
doc.add_paragraph('Function: calculateLeaseAdjustment(inputs: LeaseInputs, currentEBITDA, currentTotalDebt) -> LeaseAdjustment')
doc.add_paragraph('Inputs: totalLeasePaymentsPerYear, remainingLeaseTerm, incrementalBorrowingRate (IBR as %)')
bullets([
    'ROU Asset = PV of remaining lease payments: PMT * [(1 - (1+r)^-n) / r]',
    'Lease Liability = ROU Asset at inception',
    'EBITDA Add-Back = full annual lease payment (reclassified operating -> financing)',
    'Post-EAS48 EBITDA = current EBITDA + annual lease payment',
    'Post-EAS48 Debt = current debt + lease liability',
    'Annual Depreciation = ROU / remaining term (straight-line)',
    'Interest Expense = Lease Liability * IBR (Year 1)',
    'Edge case: If IBR = 0, ROU = payment * term (simple multiplication)',
])

doc.add_heading('7.2 EAS 31 (IAS 1) — Normalized Earnings', level=2)
doc.add_paragraph('Function: calculateNormalizedEarnings(items: NormalizationItem[], reportedNI, taxRate, shares) -> NormalizationResult')
bullets([
    'Each NormalizationItem: description, amount (positive=add-back, negative=deduction), isTaxAffected (boolean)',
    'Gross Adjustment = sum of all items',
    'Tax Effect = sum of tax-affected items * (taxRate/100)',
    'Net Adjustment = Gross - Tax Effect',
    'Normalized NI = Reported NI + Net Adjustment',
    'Normalized EPS = Normalized NI / Shares Outstanding',
])

doc.add_heading('7.3 EAS 12 (IAS 12) — Deferred Tax in EV Bridge', level=2)
doc.add_paragraph('Function: calculateDeferredTaxAdjustment(inputs: DeferredTaxInputs, baseEquityValue) -> DeferredTaxAdjustment')
bullets([
    'Net DTA = DTA - DTL',
    'If positive (DTA > DTL): adds to equity value (future tax savings)',
    'If negative (DTL > DTA): reduces equity value (future tax obligations)',
    'Adjusted Equity = Base Equity Value + Net DTA',
])

doc.add_heading('7.4 EAS 23 (IAS 33) — Basic & Diluted EPS', level=2)
doc.add_paragraph('Function: calculateEPS(netIncome, basicShares, dilutedShares, convertibleInterest=0) -> EPSResult')
bullets([
    'Basic EPS = Net Income / Weighted Average Basic Shares',
    'Raw Diluted EPS = (NI + Convertible Interest) / max(Diluted Shares, Basic Shares)',
    'Anti-Dilution Check: If rawDilutedEPS > basicEPS, securities are anti-dilutive',
    'Final Diluted EPS = anti-dilutive ? basicEPS : rawDilutedEPS',
    'Dilution % = (basicEPS - dilutedEPS) / basicEPS * 100',
    'Returns: basicEPS, dilutedEPS, basicShares, dilutedShares, dilutionPercent, isAntiDilutive',
])
doc.add_page_break()

# ═══════════════════════════════
# 8. CONFIDENCE SCORING
# ═══════════════════════════════
doc.add_heading('8. Model Confidence Scoring (0-100)', level=1)
doc.add_paragraph('File: src/utils/confidenceScore.ts (186 lines)')
doc.add_paragraph('Function: calculateConfidenceScore(financialData, assumptions, dcfResult, ddmResult, fcffVerification) -> ConfidenceScoreResult')
doc.add_paragraph('Grade: A (85+), B (70+), C (55+), D (40+), F (<40)')
add_table(['Pillar', 'Max', 'Individual Checks'], [
    ['Data Quality', '30', 'D&A > 0 (5pts), CapEx > 0 (5), Interest rate consistency (5), EBITDA > 0 (5), Revenue > 0 (5), Shares > 0 (5)'],
    ['Assumption Reasonableness', '40', 'Rev growth <= 25% (8), EBITDA margin <= 50% (8), Terminal g < GDP ~6% (8), WACC 15-35% for Egypt (8), FCFF 3-way reconciles (8)'],
    ['Model Robustness', '30', 'TV < 80% of EV (10), Value within +/-50% of market (10), DCF/DDM directional agreement (5), Interest coverage > 1.5x (5)'],
])
doc.add_paragraph('Scoring is graduated — partial points for marginal violations. Returns totalScore, grade, per-pillar scores, detailed breakdown array, and narrative explanation.')

# ═══════════════════════════════
# 9. INPUT VALIDATION
# ═══════════════════════════════
doc.add_heading('9. Input Validation Engine', level=1)
doc.add_paragraph('File: src/utils/inputValidation.ts (353 lines). Three severity tiers.')

doc.add_heading('9.1 Hard Blocks (Calculation Cannot Proceed)', level=2)
add_table(['Check', 'Condition', 'Reason'], [
    ['Revenue', '<= 0', 'Cannot perform valuation without revenue'],
    ['Shares', '<= 0', 'Cannot calculate per-share values'],
    ['WACC vs g', 'WACC <= Terminal Growth', 'Gordon Growth Model infinite/negative'],
    ['WACC', '<= 0', 'Invalid discount rate'],
    ['Projection Years', '< 1 or > 20', 'Out of reasonable range'],
    ['Stock Price', '<= 0', 'Cannot calculate upside/downside'],
    ['EBITDA Margin', '<= 0% or > 95%', 'Unrealistic operating profile'],
    ['Revenue Growth', '> 100%', 'Unreasonably aggressive'],
    ['Tax Rate', '< 0% or > 60%', 'Out of legal range'],
])

doc.add_heading('9.2 Soft Warnings (Review Recommended)', level=2)
bullets([
    'Terminal Value > 80% of EV: over-reliance on terminal assumptions',
    'Implied price deviates > 50% from market: may indicate aggressive assumptions',
    'Negative FCF in base year: DCF may extrapolate losses',
    'Interest Coverage < 1.5x: potential debt service issues',
    'D&A = 0 with positive PP&E: may understate reinvestment',
    'CapEx = 0: FCFF may be overstated',
    'Revenue growth > 25%: aggressive for most companies',
    'Debt/Equity > 3.0: highly leveraged',
])

doc.add_heading('9.3 Edge Case Auto-Adjustments', level=2)
bullets([
    'Zero interest with positive debt: impute ~2.8% of total debt',
    'COGS = 0 for banks: recalculate as Revenue - Interest Expense',
    'Negative equity: flag with warning but allow calculation',
    'Missing D&A: default to 0 with warning about understated reinvestment',
])
doc.add_page_break()

# ═══════════════════════════════
# 10. MARKET REGIONS
# ═══════════════════════════════
doc.add_heading('10. Market Region Support (USA & Egypt)', level=1)
doc.add_paragraph('File: src/constants/marketDefaults.ts (96 lines). Auto-detection: .CA suffix triggers Egypt.')

doc.add_heading('10.1 Market Defaults', level=2)
add_table(['Parameter', 'USA', 'Egypt (BUG FIXED)'], [
    ['Risk-Free Rate', '4.5% (10Y US Treasury)', '22.0% (10Y Egyptian Govt Bond) — was 27.25% CBE overnight'],
    ['Market Risk Premium', '5.5%', '5.5% (mature ERP only) — was 10% which double-counted CRP'],
    ['Terminal Growth Rate', '2.5%', '8.0% (nominal GDP growth)'],
    ['Max Terminal Growth', '4.0%', '12.0% (includes inflation)'],
    ['Default Tax Rate', '21.0%', '22.5%'],
    ['Currency', 'USD ($)', 'EGP'],
    ['Country Risk Premium', '0%', '7.5% (Damodaran, Caa1/B-) — Method B ONLY'],
    ['CBE Benchmark Rate', 'N/A', '27.25% (CBE overnight lending corridor)'],
    ['Dividend WHT', 'N/A', '10% (Article 46 bis)'],
    ['Capital Gains Tax', 'N/A', '10% (listed securities)'],
])

doc.add_heading('10.2 Egyptian Tax Categories (6 types)', level=2)
add_table(['Category', 'Rate', 'Applies To'], [
    ['Standard', '22.5%', 'Most Egyptian companies (DEFAULT)'],
    ['Industrial Zone', '0%', 'First 5 years; reverts to 22.5%'],
    ['Oil & Gas', '40.55%', 'Petroleum sector companies'],
    ['Suez Canal / EGPC / CBE', '40.0%', 'Specific state entities'],
    ['Free Zone', '0%', 'Free economic zone exports'],
    ['Custom', 'User-defined', 'Any special case'],
])

doc.add_heading('10.3 Egyptian Industry Multiples (8 sectors)', level=2)
add_table(['Sector', 'P/E', 'EV/EBITDA', 'P/S', 'P/B', 'Weights (PE/EV/PS/PB)'], [
    ['Banking', '5.5x', 'N/A', '2.0x', '1.0x', '50/0/25/25'],
    ['Real Estate', '8.0x', '7.0x', '1.5x', '0.8x', '40/35/15/10'],
    ['Telecom', '12.0x', '6.0x', '1.8x', '2.0x', '40/35/15/10'],
    ['Consumer/FMCG', '15.0x', '8.0x', '1.0x', '2.5x', '40/35/15/10'],
    ['Industrial', '7.0x', '5.0x', '0.8x', '1.2x', '40/35/15/10'],
    ['Healthcare', '18.0x', '10.0x', '2.0x', '3.0x', '40/35/15/10'],
    ['Energy', '6.0x', '4.5x', '0.6x', '1.0x', '40/35/15/10'],
    ['EGX Average', '7.0x', '5.0x', '1.2x', '1.5x', '40/35/15/10'],
])

doc.add_heading('10.4 US Industry Multiples (6 sectors)', level=2)
add_table(['Sector', 'P/E', 'EV/EBITDA', 'P/S', 'P/B'], [
    ['Technology', '28x', '18x', '6x', '8x'],
    ['Financial Services', '12x', '8x', '3x', '1.2x'],
    ['Consumer', '22x', '12x', '2x', '4x'],
    ['Healthcare', '20x', '14x', '4x', '5x'],
    ['Industrial', '18x', '10x', '1.5x', '3x'],
    ['Market Average', '20x', '12x', '3x', '4x'],
])

doc.add_heading('10.5 Egyptian Peer Groups', level=2)
doc.add_paragraph('Built-in mappings: CIB/COMI.CA -> QNBA/ADIB/FAISAL/SAIB, ETEL.CA -> VODAFONE/ORANGE, EFID.CA/Edita -> JUFO/ISPH/DOMTY, TMGH.CA/Palm Hills -> PHDC/EMFD/MNHD. Default EG group: COMI/ETEL/EFID/EAST/TMGH.')
doc.add_page_break()

# ═══════════════════════════════
# 11. API SERVICES
# ═══════════════════════════════
doc.add_heading('11. API Services', level=1)
doc.add_heading('11.1 FMP Stable API (Primary)', level=2)
doc.add_paragraph('File: src/services/stockAPI.ts. URL: https://financialmodelingprep.com/stable/')
doc.add_paragraph('Parallel fetch via Promise.all: profile, income-statement, balance-sheet-statement, cash-flow-statement, ratios-ttm.')
doc.add_paragraph('Special handling: Bank sector (COGS=0 -> Gross Profit = Rev - IntExp), long-term investment reclassification (>40% unclassified NCA), interest fallback (~2.8% of debt when API reports 0).')

doc.add_heading('11.2 Yahoo Finance (Egyptian Market)', level=2)
doc.add_paragraph('File: src/services/yahooFinanceAPI.ts. v10 quoteSummary via CORS proxy chain (allorigins -> codetabs -> corsproxy). For .CA suffixed Egyptian stocks.')

doc.add_heading('11.3 Peer Fetching', level=2)
doc.add_paragraph('Sequential with 200ms rate-limit delay. Fetches profile + 2 years IS per peer. Calculates multiples client-side from raw data.')

# ═══════════════════════════════
# 12. VALUATION STYLES
# ═══════════════════════════════
doc.add_heading('12. Valuation Styles & Scenario Presets', level=1)
doc.add_paragraph('File: src/constants/valuationStyles.ts (44 lines)')
add_table(['Parameter', 'Conservative (Value)', 'Moderate (GARP)', 'Aggressive (Growth)'], [
    ['Revenue Growth Mult', '0.7x', '1.0x', '1.4x'],
    ['WACC Adjustment', '+1.5%', '0%', '-1.0%'],
    ['Terminal Growth Mult', '0.8x', '1.0x', '1.2x'],
    ['Multiple Mult', '0.8x', '1.0x', '1.3x'],
    ['Margin Change', '-0.5%', '0%', '+1.0%'],
])
doc.add_paragraph('Applied AFTER scenario multipliers. Terminal growth capped at market-specific max (4.0% USA, 12.0% Egypt).')

# ═══════════════════════════════
# 13. STATE MANAGEMENT
# ═══════════════════════════════
doc.add_heading('13. State Management & Hooks', level=1)
add_table(['Hook', 'File', 'Responsibility'], [
    ['useFinancialData', 'hooks/useFinancialData.ts', 'Central state: financialData + assumptions + comparables. Auto-save 1s debounce. Auto-detect .CA for Egypt. Auto-calc beta/tax/WACC from API.'],
    ['useValuationCalculations', 'hooks/useValuationCalculations.ts', 'Derives ALL outputs via useMemo: WACC, FCFF, DCF, DDM, Comps, Blended, Ratios, Sensitivity, Scenarios, ReverseDCF, Confidence, chart-ready arrays.'],
    ['useHistory', 'hooks/useHistory.ts', 'Undo/Redo stack. Snapshots: {financialData, assumptions, comparables}.'],
    ['useTheme', 'hooks/useTheme.ts', 'Dark/Light mode toggle. Persists to localStorage.'],
    ['useKeyboardShortcuts', 'hooks/useKeyboardShortcuts.ts', 'Ctrl+Z (undo), Ctrl+Y (redo), Ctrl+S (save).'],
])

# ═══════════════════════════════
# 14. EXPORT
# ═══════════════════════════════
doc.add_heading('14. Export Capabilities', level=1)
add_table(['Export', 'File', 'Size', 'Key Features'], [
    ['PDF', 'pdfExport.ts', '33KB', 'Multi-page report: Cover, IS, BS, CF, Assumptions, DCF, EV Bridge, Comps, Scenarios, Disclaimer. Auto-pagination.'],
    ['Excel', 'excelExport.ts', '45KB', 'Multi-sheet: Summary, IS, BS, CF, DCF Model, Assumptions, Comparables, Quality Scorecard. Proper number formats.'],
    ['Excel Pro', 'excelExportPro.ts', '38KB', 'Live Excel formulas. Changing assumption recalculates everything. Cross-sheet references.'],
    ['JSON', 'jsonExport.ts', '5.7KB', 'CFA-grade. Engine v4.0. Three sections: metadata (9 fields), inputs (25+), calculated (WACC+DCF+DDM+multiples+sensitivity+ratios+FCFF+scenarios).'],
])
doc.add_page_break()

# ═══════════════════════════════
# 15. UI COMPONENTS
# ═══════════════════════════════
doc.add_heading('15. UI Components & Tabs', level=1)

doc.add_heading('Input Tab', level=2)
doc.add_paragraph('StockSearch (ticker + API fetch), FinancialInputForm (IS/BS/CF editable), AssumptionsPanel (WACC/growth/margin/DDM/drivers), ComparableCompanies (peer table + auto-fetch), Market Region selector, APIKeyModal.')

doc.add_heading('Valuation Tab — 15 Sub-Sections', level=2)
add_table(['#', 'Component', 'Purpose'], [
    ['1', 'ValidationAlerts.tsx', 'Market-specific warnings with severity levels and blocking indicators'],
    ['2', 'ValuationStyleSelector.tsx', 'Conservative/Moderate/Aggressive with side-by-side DCF comparison'],
    ['3', 'MarketVsFundamental.tsx', 'Market price vs intrinsic value with key statistics'],
    ['4', 'ScenarioAnalysis.tsx', 'Bear/Base/Bull scenarios + sensitivity matrix + Monte Carlo'],
    ['5', 'ValuationSummaryCards.tsx', 'DCF, DDM, Comps, Blended value cards with upside'],
    ['6', 'BlendedWeightSlider.tsx', 'Adjustable DCF/Comps weight slider with live preview'],
    ['7', 'ComparableBreakdown.tsx', 'Individual method implied prices (P/E, EV/EBITDA, P/S, P/B)'],
    ['8', 'BaseYearFCF.tsx', 'Base year free cash flow breakdown'],
    ['9', 'DCFProjectionsTable.tsx', 'Year-by-year: Revenue, EBITDA, D&A, EBIT, NOPAT, CapEx, WC, FCFF, DF, PV'],
    ['10', 'DDMValuation.tsx', 'Gordon Growth, Two-Stage, H-Model results with applicability check'],
    ['11', 'FCFFReconciliation.tsx', 'Three-method FCFF cross-verification with match/mismatch indicators'],
    ['12', 'KeyMetricsGrid.tsx', '30+ financial ratios in card format, organized by category'],
    ['13', 'ReverseDCFSection.tsx', 'Market-implied growth rate analysis with narrative'],
    ['14', 'QualityScorecard.tsx', '4-category quality assessment with letter grade and premium/discount'],
    ['15', 'EASComplianceSection.tsx', 'EAS 48/31/12/23 compliance panels with live calculations'],
])

doc.add_heading('Charts Tab', level=2)
doc.add_paragraph('Revenue/FCF/EBITDA projections (line/area), Valuation comparison bars, Football field chart (range per method). Interactive Recharts with currency-aware tooltips.')

doc.add_heading('Shared & Other Components', level=2)
doc.add_paragraph('ErrorBoundary, LoadingFallback, InputField, CalculationAuditTrail, DebouncedInput, Tooltip, AIReport (AI investment thesis), QuickStartGuide, HowToBuildModal, UserAuth, WolfLogo, Header, CompanyHeader, TabNavigation, Footer.')

# ═══════════════════════════════
# 16-18. UTILITIES, WORKER, CONFIG
# ═══════════════════════════════
doc.add_heading('16. Formatting & Utility Functions', level=1)
doc.add_paragraph('File: src/utils/formatters.ts (6KB). Currency-aware formatting:')
add_table(['Function', 'USD Example', 'EGP Example'], [
    ['formatCurrency()', '$1,500,000', '1,500,000 EGP'],
    ['formatCurrencyShort()', '$1.50B', '1.50B EGP'],
    ['formatPercent()', '7.50%', '7.50%'],
    ['formatMultiple()', '15.5x', '15.5x'],
    ['formatMillions()', '$1,500.00M', '1,500.00M EGP'],
    ['formatPrice()', '$45.50', '45.50 EGP'],
    ['formatShares()', '500.00M', '500.00M'],
])

doc.add_heading('17. Web Worker', level=1)
doc.add_paragraph('File: src/workers/monteCarlo.worker.ts. Dedicated thread. Receives financialData + assumptions via postMessage. Runs 5,000 simulations. Returns statistical distribution. Prevents UI blocking.')

doc.add_heading('18. Configuration & Build', level=1)
bullets([
    'Vite config: React plugin + Tailwind CSS plugin + single-file plugin for deployment',
    'Dev server: npm run dev (Vite HMR)',
    'Production build: npm run build (outputs to dist/)',
    'Environment: .env.example defines VITE_FMP_API_KEY',
    'TypeScript: Strict mode, ES2020 target, React JSX transform',
    'Single-file build: vite-plugin-singlefile inlines all assets for portable deployment',
])
doc.add_page_break()

# ═══════════════════════════════
# 19. KNOWN ISSUES & FUTURE
# ═══════════════════════════════
doc.add_heading('19. Known Issues & Future Development', level=1)
doc.add_heading('Current Known Issues', level=2)
bullets([
    'Legacy valuation.ts has duplicate calculation logic — should migrate to valuationEngine.ts single source',
    'Monte Carlo worker duplicates logic from advancedAnalysis.ts — consider shared module',
    'localStorage auto-save validates JSON but does not auto-restore on mount',
    'Piotroski F-Score is stubbed (returns 0) — needs 9-point implementation',
    'Forward P/E uses simple 15% growth proxy — should use projected Year 1 EPS',
])
doc.add_heading('Suggested Future Enhancements', level=2)
bullets([
    'Multi-year historical trend analysis (currently single year)',
    'Custom peer group definition with saved presets',
    'Automated regression for beta from price history',
    'Full Piotroski F-Score (9-criteria) implementation',
    'SOTP (Sum of the Parts) valuation for conglomerates',
    'Real-time price updates via WebSocket',
    'Cloud-based valuation storage with user auth',
    'Waterfall chart for EV-to-equity bridge visualization',
    'Additional emerging markets (Saudi KSA, UAE, Nigeria)',
    'Unit and integration test suite (Jest/Vitest)',
    'Year-by-year varying projection drivers (useConstantDrivers = false)',
    'Mobile-responsive layout optimizations',
    'Automated PDF/Excel scheduling and email delivery',
])

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'WOLF_Engine_Complete_Overview.docx')
doc.save(output_path)
sz = os.path.getsize(output_path)
print(f'Document saved to: {output_path}')
print(f'Size: {sz:,} bytes ({sz/1024:.1f} KB)')
