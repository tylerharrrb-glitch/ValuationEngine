from docx import Document
import os

doc = Document(r'C:\Users\mcdr\Desktop\WOLF_Agent_Prompt.docx')
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'WOLF_Agent_Prompt.txt')

with open(output_path, 'w', encoding='utf-8') as f:
    for para in doc.paragraphs:
        f.write(para.text + '\n')
    
    # Also check tables
    for i, table in enumerate(doc.tables):
        f.write(f'\n--- TABLE {i+1} ---\n')
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            f.write(' | '.join(cells) + '\n')

print(f"Written to {output_path}")
