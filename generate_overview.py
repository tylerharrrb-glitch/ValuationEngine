"""
Generate WOLF Valuation Engine - Complete Overview Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()

# ════════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('WOLF Valuation Engine')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = subtitle.add_run('Complete Engine Overview & Developer Reference')
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run('Version 1.0  |  February 2026')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = desc.add_run(
    'Professional-grade equity valuation platform built with React, TypeScript, and Vite.\n'
    'Supports DCF, Comparable Company Analysis, Monte Carlo Simulation, and more.\n'
    'Dual-market support: USA (NYSE/NASDAQ) and Egypt (EGX).'
)
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. Technology Stack & Architecture',
    '3. Project Structure & File Map',
    '4. Data Models & TypeScript Interfaces',
    '5. Core Valuation Engine',
    '   5.1 WACC Calculation',
    '   5.2 DCF Valuation',
    '   5.3 Comparable Company Valuation',
    '   5.4 Blended Valuation & Recommendations',
    '   5.5 Sensitivity Analysis',
    '6. Advanced Analysis Module',
    '   6.1 Reverse DCF',
    '   6.2 Monte Carlo Simulation',
    '   6.3 Sector Benchmarking',
    '   6.4 Quality Scorecard',
    '7. Market Region Support (USA & Egypt)',
    '8. API Services (FMP & Yahoo Finance)',
    '9. Valuation Styles & Scenario Analysis',
    '10. State Management & History',
    '11. Export Capabilities (PDF & Excel)',
    '12. UI Components & Tabs',
    '13. Formatting & Utility Functions',
    '14. Web Worker (Monte Carlo)',
    '15. Configuration & Build',
    '16. Known Issues & Future Development Areas',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'WOLF Valuation Engine is a professional-grade, browser-based equity valuation platform '
    'that combines multiple valuation methodologies into a single, interactive tool. It is designed '
    'for financial analysts, portfolio managers, and investment professionals who need to perform '
    'rigorous fundamental analysis on publicly traded companies.'
)
doc.add_paragraph(
    'The engine supports two markets: the United States (NYSE/NASDAQ) and Egypt (EGX), with '
    'market-specific defaults for risk-free rates, equity risk premiums, tax rates, and industry '
    'multiples. Live data is fetched from Financial Modeling Prep (FMP) API and Yahoo Finance.'
)
doc.add_heading('Key Capabilities', level=3)
bullets = [
    'Discounted Cash Flow (DCF) valuation with mid-year convention',
    'Comparable Company Analysis using 4 methods (P/E, EV/EBITDA, P/S, P/B)',
    'Blended valuation with adjustable DCF/Comps weighting',
    'Scenario analysis (Bull/Bear/Base cases)',
    'Three valuation styles: Conservative, Moderate, Aggressive',
    'Sensitivity analysis matrix (WACC vs Terminal Growth)',
    'Reverse DCF (implied growth rate extraction)',
    'Monte Carlo simulation (5,000 iterations via Web Worker)',
    'Sector benchmarking against industry averages',
    'Quality scorecard (Economic Moat, Financial Health, Growth, Capital Allocation)',
    'AI-powered investment report generation',
    'Export to PDF and Excel (with live formulas)',
    'Real-time stock data via FMP Stable API',
    'Egyptian market support via Yahoo Finance fallback with CORS proxy chain',
    'Undo/Redo history with auto-save to localStorage',
    'Dark/Light theme, keyboard shortcuts, lazy-loaded tabs',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ════════════════════════════════════════════════════════════════
doc.add_heading('2. Technology Stack & Architecture', level=1)

doc.add_heading('Frontend Framework', level=2)
add_table(
    ['Technology', 'Version', 'Purpose'],
    [
        ['React', '19.2.3', 'UI framework with hooks-based components'],
        ['TypeScript', '5.9.3', 'Static typing for all source files'],
        ['Vite', '7.2.4', 'Build tool and dev server (HMR)'],
        ['Tailwind CSS', '4.1.17', 'Utility-first CSS framework'],
    ]
)

doc.add_heading('Key Libraries', level=2)
add_table(
    ['Library', 'Version', 'Purpose'],
    [
        ['Recharts', '3.7.0', 'Interactive charts (bar, line, area)'],
        ['Lucide React', '0.563.0', 'Icon library (feather-style)'],
        ['jsPDF + autoTable', '4.1.0 / 5.0.7', 'PDF report generation'],
        ['SheetJS (xlsx)', '0.18.5', 'Excel workbook export with formulas'],
        ['clsx + tailwind-merge', '2.1.1 / 3.4.0', 'Conditional class merging'],
        ['vite-plugin-singlefile', '2.3.0', 'Single HTML file build output'],
    ]
)

doc.add_heading('Architecture Pattern', level=2)
doc.add_paragraph(
    'The application follows a hooks-based architecture with clear separation of concerns:'
)
arch_items = [
    'App.tsx — Thin orchestrator that composes hooks and layout components.',
    'Custom Hooks (useFinancialData, useValuationCalculations, useTheme, useHistory, useKeyboardShortcuts) — '
    'All business logic and state management lives in hooks. Calculations are memoized with useMemo.',
    'Pure Utility Functions (utils/) — All financial calculations are pure functions with no React dependencies. '
    'This makes them testable and reusable.',
    'Services Layer (services/) — API integration with FMP and Yahoo Finance, including error handling and data mapping.',
    'Component Layer — Presentational components organized by feature: layout, input, valuation sections, charts, shared.',
    'Lazy Loading — Tab content (Input, Valuation, Charts) is lazy-loaded with React.lazy() and Suspense.',
    'Error Boundaries — Wrap each tab and the root app for graceful error recovery.',
    'Web Worker — Monte Carlo simulation runs in a dedicated worker thread to avoid blocking the UI.',
]
for item in arch_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 3. PROJECT STRUCTURE
# ════════════════════════════════════════════════════════════════
doc.add_heading('3. Project Structure & File Map', level=1)

doc.add_paragraph('The project root contains configuration files and the src/ directory with all source code:')

structure = """ValuationEngine/
├── index.html                    # SPA entry point
├── package.json                  # Dependencies & scripts
├── vite.config.ts                # Vite build configuration
├── tsconfig.json                 # TypeScript configuration
├── .env.example                  # Environment variable template
│
└── src/
    ├── main.tsx                  # React DOM bootstrap with ErrorBoundary
    ├── App.tsx                   # Main orchestrator component
    ├── index.css                 # Root CSS imports
    ├── vite-env.d.ts             # Vite type declarations
    │
    ├── types/
    │   └── financial.ts          # All TypeScript interfaces & types
    │
    ├── constants/
    │   ├── initialData.ts        # Default financial data & assumptions
    │   ├── marketDefaults.ts     # USA/Egypt market parameters
    │   └── valuationStyles.ts    # Conservative/Moderate/Aggressive presets
    │
    ├── data/
    │   └── sampleData.ts         # Sample demo data
    │
    ├── hooks/
    │   ├── useFinancialData.ts   # State mgmt with history & auto-save
    │   ├── useValuationCalculations.ts  # Memoized valuation computations
    │   ├── useHistory.ts         # Undo/Redo stack management
    │   ├── useTheme.ts           # Dark/Light mode toggle
    │   └── useKeyboardShortcuts.ts  # Ctrl+Z, Ctrl+Y, Ctrl+S
    │
    ├── services/
    │   ├── stockAPI.ts           # FMP Stable API (primary data source)
    │   └── yahooFinanceAPI.ts    # Yahoo Finance fallback (Egyptian stocks)
    │
    ├── utils/
    │   ├── valuationEngine.ts    # Core engine: WACC, DCF, comps, sensitivity
    │   ├── valuation.ts          # Legacy calc helpers
    │   ├── advancedAnalysis.ts   # Reverse DCF, Monte Carlo, benchmarking, scorecard
    │   ├── formatters.ts         # Currency, number, percent formatting
    │   ├── industryMapping.ts    # Ticker-to-industry sector mapping
    │   ├── pdfExport.ts          # PDF report generation
    │   ├── excelExport.ts        # Excel workbook export
    │   ├── excelExportPro.ts     # Excel export with live formulas
    │   ├── cn.ts                 # Class name utility
    │   └── calculations/
    │       ├── dcf.ts            # DCF projection & valuation functions
    │       ├── comparables.ts    # Comparable company calc functions
    │       ├── metrics.ts        # Key financial metrics & recommendations
    │       └── scenarios.ts      # Bull/Bear/Base scenario functions
    │
    ├── workers/
    │   └── monteCarlo.worker.ts  # Web Worker for Monte Carlo simulation
    │
    └── components/
        ├── layout/
        │   ├── Header.tsx        # Top bar with controls
        │   ├── CompanyHeader.tsx  # Company name, price, recommendation
        │   ├── TabNavigation.tsx  # Input/Valuation/Charts tabs
        │   └── Footer.tsx        # App footer
        ├── shared/
        │   ├── ErrorBoundary.tsx  # React error boundary
        │   ├── LoadingFallback.tsx # Lazy load spinner
        │   ├── InputField.tsx    # Reusable form input
        │   └── CalculationAuditTrail.tsx # History status bar
        ├── input/
        │   └── InputTab.tsx      # All data entry forms
        ├── charts/
        │   └── ChartsTab.tsx     # Revenue projection, football field, etc.
        ├── valuation/
        │   ├── ValuationTab.tsx  # Valuation tab orchestrator
        │   └── sections/        # 12 valuation sub-sections
        │       ├── ValidationAlerts.tsx
        │       ├── ValuationStyleSelector.tsx
        │       ├── MarketVsFundamental.tsx
        │       ├── ScenarioAnalysis.tsx
        │       ├── ValuationSummaryCards.tsx
        │       ├── BlendedWeightSlider.tsx
        │       ├── ComparableBreakdown.tsx
        │       ├── BaseYearFCF.tsx
        │       ├── DCFProjectionsTable.tsx
        │       ├── KeyMetricsGrid.tsx
        │       ├── ReverseDCFSection.tsx
        │       └── QualityScorecard.tsx
        ├── StockSearch.tsx       # Ticker search & API fetch
        ├── FinancialInputForm.tsx # Income/BS/CF input forms
        ├── AssumptionsPanel.tsx  # WACC & growth assumptions
        ├── ComparableCompanies.tsx # Peer company table
        ├── DCFResults.tsx        # DCF results display
        ├── KeyMetrics.tsx        # Key metrics cards
        ├── ValuationSummary.tsx  # Summary with football field
        ├── FootballFieldChart.tsx # Football field visualization
        ├── ScenarioToggle.tsx    # Bear/Base/Bull selector
        ├── DebouncedInput.tsx    # Debounced numeric input
        ├── Tooltip.tsx           # Hover tooltip component
        ├── AIReport.tsx          # AI investment report generator
        ├── QuickStartGuide.tsx   # Onboarding modal
        ├── HowToBuildModal.tsx   # Educational modal
        ├── APIKeyModal.tsx       # API key configuration
        ├── UserAuth.tsx          # User authentication
        └── WolfLogo.tsx          # WOLF brand logo SVG"""

p = doc.add_paragraph()
r = p.add_run(structure)
r.font.name = 'Consolas'
r.font.size = Pt(7)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 4. DATA MODELS
# ════════════════════════════════════════════════════════════════
doc.add_heading('4. Data Models & TypeScript Interfaces', level=1)
doc.add_paragraph('All types are defined in src/types/financial.ts. The core data structures are:')

doc.add_heading('FinancialData', level=2)
doc.add_paragraph('The central data object that holds all company financial information:')
add_table(
    ['Field', 'Type', 'Description'],
    [
        ['companyName', 'string', 'Company display name'],
        ['ticker', 'string', 'Stock ticker symbol (e.g., AAPL, COMI.CA)'],
        ['sharesOutstanding', 'number', 'Total shares outstanding'],
        ['currentStockPrice', 'number', 'Current market price per share'],
        ['lastReportedDate?', 'string', 'Last fiscal year end date from API'],
        ['sector?', 'string', 'Company sector (e.g., "Technology", "Financial Services")'],
        ['incomeStatement', 'IncomeStatement', 'Revenue, COGS, OpEx, NetIncome, D&A, etc.'],
        ['balanceSheet', 'BalanceSheet', 'Assets, Liabilities, Equity (19 fields)'],
        ['cashFlowStatement', 'CashFlowStatement', 'Operating CF, CapEx, FCF, Dividends'],
    ]
)

doc.add_heading('IncomeStatement', level=2)
add_table(
    ['Field', 'Description'],
    [
        ['revenue', 'Total revenue'],
        ['costOfGoodsSold', 'Cost of goods sold (Interest Expense for banks)'],
        ['grossProfit', 'Revenue minus COGS'],
        ['operatingExpenses', 'SG&A and other operating expenses'],
        ['operatingIncome', 'EBIT (Earnings before interest and taxes)'],
        ['interestExpense', 'Interest on debt'],
        ['taxExpense', 'Income tax expense'],
        ['netIncome', 'Bottom-line earnings'],
        ['depreciation', 'Depreciation expense'],
        ['amortization', 'Amortization expense'],
    ]
)

doc.add_heading('BalanceSheet', level=2)
doc.add_paragraph(
    'Contains 19 fields organized into Current Assets (cash, marketableSecurities, accountsReceivable, '
    'inventory, otherCurrentAssets, totalCurrentAssets), Non-Current Assets (propertyPlantEquipment, '
    'longTermInvestments, goodwill, intangibleAssets, otherNonCurrentAssets, totalAssets), '
    'Current Liabilities (accountsPayable, shortTermDebt, otherCurrentLiabilities, totalCurrentLiabilities), '
    'Non-Current Liabilities (longTermDebt, otherNonCurrentLiabilities, totalLiabilities), and Equity (totalEquity).'
)

doc.add_heading('CashFlowStatement', level=2)
add_table(
    ['Field', 'Description'],
    [
        ['operatingCashFlow', 'Cash from operations'],
        ['capitalExpenditures', 'Capital expenditures (stored as positive)'],
        ['freeCashFlow', 'FCF = Operating CF - CapEx'],
        ['dividendsPaid', 'Cash dividends paid (stored as positive)'],
        ['netChangeInCash', 'Net change in cash position'],
    ]
)

doc.add_heading('ValuationAssumptions', level=2)
add_table(
    ['Field', 'Default', 'Description'],
    [
        ['discountRate', '10%', 'WACC / discount rate for DCF'],
        ['terminalGrowthRate', '2.5%', 'Perpetual growth rate (Gordon Growth)'],
        ['projectionYears', '5', 'Number of DCF projection years'],
        ['revenueGrowthRate', '8%', 'Annual revenue growth rate'],
        ['marginImprovement', '0.5%', 'Annual FCF margin improvement'],
        ['taxRate', '21%', 'Corporate tax rate'],
        ['riskFreeRate', '4.5%', 'Risk-free rate (10Y Treasury)'],
        ['marketRiskPremium', '5.5%', 'Equity risk premium'],
        ['beta', '1.1', 'Company beta coefficient'],
    ]
)

doc.add_heading('Other Key Types', level=2)
other_types = [
    'ComparableCompany — Peer company with name, ticker, P/E, EV/EBITDA, P/S, P/B, marketCap, revenue.',
    'DCFProjection — Per-year projection: year, revenue, ebitda, freeCashFlow, discountFactor, presentValue.',
    'ValuationResult — Method name, total value, per-share value, upside percentage.',
    'MarketRegion — Union type: "USA" | "Egypt".',
    'HistoryState — Snapshot of financialData + assumptions + comparables for undo/redo.',
    'ThemeClasses — CSS class strings for dark/light mode theming.',
    'ScenarioType — Union type: "bear" | "base" | "bull".',
    'ValuationStyleKey — Union type: "conservative" | "moderate" | "aggressive".',
]
for item in other_types:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 5. CORE VALUATION ENGINE
# ════════════════════════════════════════════════════════════════
doc.add_heading('5. Core Valuation Engine', level=1)
doc.add_paragraph(
    'The core valuation logic is split across two layers: valuationEngine.ts (comprehensive, '
    'documented engine with WACC, DCF, Comps, Blended, Sensitivity, and Financial Metrics) '
    'and utils/calculations/ (refactored pure functions used by the React hooks).'
)

doc.add_heading('5.1 WACC Calculation', level=2)
doc.add_paragraph('File: src/utils/valuationEngine.ts — calculateWACC()')
doc.add_paragraph('Formula: WACC = (E/V) × Re + (D/V) × Rd × (1 - Tc)')
wacc_items = [
    'Step 1: Cost of Equity via CAPM — Re = Rf + β × (Rm - Rf)',
    'Step 2: After-Tax Cost of Debt — Rd × (1 - Tc)',
    'Step 3: Capital Structure Weights — E/V and D/V from market cap and total debt',
    'Step 4: Weighted average of cost of equity and after-tax cost of debt',
]
for item in wacc_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.2 DCF Valuation', level=2)
doc.add_paragraph('Files: src/utils/valuationEngine.ts & src/utils/calculations/dcf.ts')
dcf_items = [
    'Projects revenue forward using compounding growth rate for N years (default 5)',
    'Calculates FCF margin = base FCF/Revenue + incremental margin improvement per year',
    'EBITDA margin tracks independently for display purposes',
    'Uses mid-year convention for discounting (year - 0.5) in the engine; end-of-year in the calculations module',
    'Terminal Value via Gordon Growth Model: TV = FCF(n) × (1+g) / (WACC - g)',
    'Fallback to exit multiple (15× FCF) if WACC ≤ terminal growth rate',
    'Enterprise Value = Sum of PV(FCFs) + PV(Terminal Value)',
    'Equity Value = EV - Total Debt + Cash',
    'Implied Share Price = Equity Value / Shares Outstanding',
    'Upside = (Implied - Current) / Current × 100%',
]
for item in dcf_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('5.3 Comparable Company Valuation', level=2)
doc.add_paragraph('Files: src/utils/valuationEngine.ts & src/utils/calculations/comparables.ts')
doc.add_paragraph('Four methods calculate implied share price from peer multiples:')
add_table(
    ['Method', 'Formula', 'Weight (Standard)', 'Weight (Banks)'],
    [
        ['P/E', 'EPS × Median Peer P/E', '40%', '60%'],
        ['EV/EBITDA', '(EBITDA × Multiple - Debt + Cash) / Shares', '35%', '0% (not meaningful for banks)'],
        ['P/S', 'Revenue per Share × Median Peer P/S', '15%', '0%'],
        ['P/B', 'Book Value per Share × Median Peer P/B', '10%', '40%'],
    ]
)
doc.add_paragraph(
    'If no user comparables are provided, industry default multiples are used based on the ticker\'s '
    'sector mapping. Multiples are further adjusted by the valuation style multiplier (0.8× conservative, '
    '1.0× moderate, 1.3× aggressive).'
)

doc.add_heading('5.4 Blended Valuation & Recommendations', level=2)
doc.add_paragraph(
    'The blended value combines DCF and Comparable values using an adjustable weight slider '
    '(default 60% DCF / 40% Comps). If no valid comparables exist, it falls back to 100% DCF.'
)
doc.add_paragraph('Recommendation thresholds based on upside:')
add_table(
    ['Upside Range', 'Recommendation'],
    [
        ['> +20%', 'STRONG BUY'],
        ['+10% to +20%', 'BUY'],
        ['-10% to +10%', 'HOLD'],
        ['-20% to -10%', 'SELL'],
        ['< -20%', 'STRONG SELL'],
    ]
)

doc.add_heading('5.5 Sensitivity Analysis', level=2)
doc.add_paragraph(
    'Generates a 5×5 matrix of implied share prices for combinations of WACC (base ± 2%) and '
    'Terminal Growth Rate (base ± 1%). Each cell recalculates the full DCF. Invalid combinations '
    '(WACC ≤ growth) show $0.'
)

doc.add_heading('5.6 Financial Metrics (30+ Ratios)', level=2)
doc.add_paragraph('The calculateFinancialMetrics() function computes:')
add_table(
    ['Category', 'Metrics'],
    [
        ['Profitability', 'Gross Margin, Operating Margin, Net Margin, EBITDA Margin'],
        ['Returns', 'ROE, ROA, ROIC'],
        ['Valuation', 'P/E, EV/EBITDA, EV/Sales, P/B, P/S'],
        ['Liquidity', 'Current Ratio, Quick Ratio'],
        ['Leverage', 'Debt/Equity, Debt/Assets, Interest Coverage'],
        ['Cash Flow', 'FCF Yield, FCF Margin, Operating CF / Debt'],
    ]
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 6. ADVANCED ANALYSIS MODULE
# ════════════════════════════════════════════════════════════════
doc.add_heading('6. Advanced Analysis Module', level=1)
doc.add_paragraph('File: src/utils/advancedAnalysis.ts — Four advanced analytical tools:')

doc.add_heading('6.1 Reverse DCF', level=2)
doc.add_paragraph(
    'Answers: "What growth rate is the market pricing in?" Uses binary search (100 iterations) '
    'to find the revenue growth rate that produces the current stock price as the DCF output. '
    'Compares implied growth to the user\'s base case assumption and categorizes market expectations '
    'as "aggressive", "reasonable", or "conservative" with a narrative explanation.'
)

doc.add_heading('6.2 Monte Carlo Simulation', level=2)
doc.add_paragraph(
    'Runs 5,000 DCF simulations with randomized parameters using Gaussian (normal) distributions:'
)
add_table(
    ['Parameter', 'Base Value', 'Standard Deviation'],
    [
        ['Revenue Growth', 'User assumption', '± 4%'],
        ['WACC', 'User assumption', '± 1.5%'],
        ['Terminal Growth', 'User assumption', '± 0.8%'],
        ['Margin Improvement', 'User assumption', '± 0.5%'],
    ]
)
doc.add_paragraph(
    'Outputs: mean, median, std dev, percentiles (5th/25th/75th/95th), probability of price exceeding '
    'current price and base case, and a 20-bucket histogram distribution. Also available as a Web Worker '
    '(src/workers/monteCarlo.worker.ts) for non-blocking execution.'
)

doc.add_heading('6.3 Sector Benchmarking', level=2)
doc.add_paragraph(
    'Compares the company against sector averages (TECH, FINANCE, DEFAULT) across 10 metrics: '
    'Gross Margin, Operating Margin, Net Margin, ROE, ROA, FCF Margin, Debt/Equity, Current Ratio, '
    'P/E, EV/EBITDA. Each metric is scored as top/above/average/below/bottom. Overall score (0-100) '
    'generates a rating (Excellent, Above Average, Average, Below Average) with an insight narrative.'
)

doc.add_heading('6.4 Quality Scorecard', level=2)
doc.add_paragraph('Scores companies across 4 categories (each out of 10 points):')
add_table(
    ['Category', 'Factors Evaluated'],
    [
        ['Economic Moat', 'Gross Margin, Operating Margin, ROE, Market Position (revenue size)'],
        ['Financial Health', 'Current Ratio, Debt/Equity, Interest Coverage, Cash/Debt ratio'],
        ['Growth & Profitability', 'Net Margin, FCF Generation, Earnings Quality (OCF > NI)'],
        ['Capital Allocation', 'CapEx/Revenue, Dividend Payout, FCF Conversion, Cash Accumulation'],
    ]
)
doc.add_paragraph(
    'Total score maps to a letter grade (A+ through F) with a quality premium/discount adjustment: '
    '+15% (A+), +10% (A), +5% (B+), 0% (B), -5% (C), -10% (D), -15% (F).'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 7. MARKET REGION SUPPORT
# ════════════════════════════════════════════════════════════════
doc.add_heading('7. Market Region Support (USA & Egypt)', level=1)
doc.add_paragraph('The engine supports dual-market operation via MarketRegion type and MARKET_DEFAULTS:')

add_table(
    ['Parameter', 'USA', 'Egypt'],
    [
        ['Risk-Free Rate', '4.5% (10Y Treasury)', '27.25% (CBE Rate)'],
        ['Market Risk Premium', '5.5%', '10.0%'],
        ['Terminal Growth Rate', '2.5%', '5.0%'],
        ['Max Terminal Growth', '4.0%', '5.0%'],
        ['Default Tax Rate', '21.0%', '22.5%'],
        ['Currency', 'USD ($)', 'EGP'],
        ['Country Risk Premium', '0%', '4.5%'],
    ]
)

doc.add_heading('Egyptian Industry Multiples', level=2)
add_table(
    ['Sector', 'P/E', 'EV/EBITDA', 'P/S', 'P/B'],
    [
        ['Banking', '5.5×', '4.0×', '2.0×', '1.2×'],
        ['Telecom', '8.0×', '4.5×', '1.5×', '1.8×'],
        ['Consumer/Food', '12.0×', '7.0×', '1.2×', '2.5×'],
        ['Real Estate', '6.0×', '8.0×', '2.0×', '0.8×'],
        ['Market Average', '7.0×', '5.0×', '1.2×', '1.5×'],
    ]
)

doc.add_paragraph(
    'Auto-detection: When a ticker ends with .CA suffix, the engine automatically switches to Egypt '
    'market region and adjusts all WACC components, tax rate, and multiples.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 8. API SERVICES
# ════════════════════════════════════════════════════════════════
doc.add_heading('8. API Services', level=1)

doc.add_heading('8.1 Financial Modeling Prep (FMP) — Primary', level=2)
doc.add_paragraph('File: src/services/stockAPI.ts')
doc.add_paragraph(
    'Uses FMP\'s 2026 Stable API format (https://financialmodelingprep.com/stable/). '
    'Fetches all data in parallel via Promise.all for performance:'
)
fmp_items = [
    'profile — Company name, price, market cap, beta, shares outstanding, sector',
    'income-statement — Revenue, COGS, operating income, net income, D&A, EBITDA',
    'balance-sheet-statement — All asset, liability, and equity line items',
    'cash-flow-statement — Operating CF, CapEx, FCF, dividends, net change in cash',
    'ratios-ttm — P/E, EV/EBITDA, P/S, P/B for comparable company fetching',
]
for item in fmp_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph(
    'Special handling: Bank sector detection (COGS=0, recalculate Gross Profit as Revenue - Interest Expense). '
    'Long-term investment reclassification heuristic for companies like Apple where >40% of non-current assets are unclassified. '
    'Interest expense fallback (~2.8% of total debt) when API reports 0 but debt exists.'
)

doc.add_heading('8.2 Yahoo Finance — Egyptian Market Fallback', level=2)
doc.add_paragraph('File: src/services/yahooFinanceAPI.ts')
doc.add_paragraph(
    'Uses Yahoo Finance v10 quoteSummary endpoint via CORS proxy chain (3 proxies tried in order): '
    'allorigins → codetabs → corsproxy. Fetches price, defaultKeyStatistics, incomeStatementHistory, '
    'balanceSheetHistory, and cashflowStatementHistory modules. Applied specially for .CA suffixed Egyptian stocks.'
)

doc.add_heading('8.3 Peer Company Fetching', level=2)
doc.add_paragraph(
    'Built-in peer groups for major US stocks (AAPL→MSFT/GOOGL/AMZN/META, etc.) and Egyptian sectors '
    '(COMI.CA→QNBA/ADIB/FAISAL, etc.). fetchAllPeerData() fetches profile + 2 years of income statements '
    'sequentially with 200ms delay for rate-limit compliance. Calculates multiples client-side from raw data.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 9. VALUATION STYLES & SCENARIOS
# ════════════════════════════════════════════════════════════════
doc.add_heading('9. Valuation Styles & Scenario Analysis', level=1)

doc.add_heading('Valuation Styles', level=2)
add_table(
    ['Parameter', 'Conservative', 'Moderate', 'Aggressive'],
    [
        ['Revenue Growth Mult', '0.7×', '1.0×', '1.4×'],
        ['WACC Adjustment', '+1.5%', '0%', '-1.0%'],
        ['Terminal Growth Mult', '0.8×', '1.0×', '1.2×'],
        ['Multiple Mult', '0.8×', '1.0×', '1.3×'],
        ['Margin Change', '-0.5%', '0%', '+1.0%'],
    ]
)

doc.add_heading('Scenario Multipliers', level=2)
doc.add_paragraph(
    'Three scenarios (Bear, Base, Bull) further adjust assumptions. The ScenarioToggle component '
    'defines multipliers for revenueGrowth, wacc, terminalGrowth, and marginChange. '
    'These are applied BEFORE valuation style adjustments.'
)
doc.add_paragraph(
    'Scenario Cases in the calculations module: Bear = 0.4× growth, +2.5% WACC, 0.6× terminal, -1.5% margin. '
    'Bull = 2.0× growth, -2.5% WACC, 1.5× terminal, +2.5% margin.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 10. STATE MANAGEMENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('10. State Management & History', level=1)
doc.add_paragraph('File: src/hooks/useFinancialData.ts')

state_items = [
    'Central hook manages financialData, assumptions, and comparables state.',
    'All state updates flow through history-integrated setters that auto-snapshot to undo/redo stack.',
    'useHistory hook maintains an array of HistoryState snapshots with a pointer for undo/redo.',
    'Auto-save to localStorage every 1 second (debounced) — persists across browser sessions.',
    'handleStockDataFetched() auto-detects Egyptian stocks by .CA suffix and switches market region.',
    'Auto-calculates Beta, Tax Rate, and WACC from fetched API data.',
    'handleLoadValuation() restores a complete saved state (all three pieces + history snapshot).',
]
for item in state_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('useValuationCalculations Hook', level=2)
doc.add_paragraph(
    'File: src/hooks/useValuationCalculations.ts — Derives ALL valuation outputs. Uses useMemo extensively. '
    'Adjusts assumptions by applying scenario multipliers THEN valuation style adjustments. '
    'Terminal growth is capped at sustainable rates (2.5% USA, 5.0% Egypt). '
    'Computes: adjustedAssumptions, dcfProjections, dcfValue, scenarioCases, industryMultiples, '
    'comparableValuations, blendedValue, upside, keyMetrics, recommendation, footballFieldData, '
    'revenueProjectionData, and valuationComparisonData (chart-ready arrays).'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 11. EXPORT CAPABILITIES
# ════════════════════════════════════════════════════════════════
doc.add_heading('11. Export Capabilities', level=1)

doc.add_heading('11.1 PDF Export', level=2)
doc.add_paragraph(
    'File: src/utils/pdfExport.ts — Uses jsPDF + jspdf-autotable. '
    'Generates multi-page research report: Cover with company info & recommendation, '
    'Income Statement, Balance Sheet, Cash Flow Statement, Assumptions & WACC, '
    'DCF Projections table, DCF Summary (EV bridge), Comparable Valuations, Scenario Analysis, '
    'and Disclaimer. Auto-pagination with header orphan prevention.'
)

doc.add_heading('11.2 Excel Export (Standard)', level=2)
doc.add_paragraph(
    'File: src/utils/excelExport.ts — Uses SheetJS (xlsx). Creates multi-sheet workbook: '
    'Summary, Income Statement, Balance Sheet, Cash Flow, DCF Model, Assumptions, Comparables, '
    'Quality Scorecard. Numeric cells with proper Excel number formats ($#,##0.00, 0.00%, etc.).'
)

doc.add_heading('11.3 Excel Export Pro (with Formulas)', level=2)
doc.add_paragraph(
    'File: src/utils/excelExportPro.ts — Advanced Excel export that writes live Excel formulas '
    'instead of static values. The exported workbook is fully editable — changing an assumption '
    'recalculates everything. Supports formula references between cells and sheets.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 12. UI COMPONENTS
# ════════════════════════════════════════════════════════════════
doc.add_heading('12. UI Components & Tabs', level=1)

doc.add_heading('Three Main Tabs', level=2)

doc.add_heading('Input Tab', level=3)
doc.add_paragraph(
    'StockSearch (ticker lookup + API fetch), FinancialInputForm (Income Statement, Balance Sheet, '
    'Cash Flow editable tables), AssumptionsPanel (WACC components, growth, margin), '
    'ComparableCompanies (peer table with auto-fetch), Market Region selector (USA/Egypt), '
    'APIKeyModal for key management.'
)

doc.add_heading('Valuation Tab', level=3)
doc.add_paragraph('12 sub-sections in order:')
val_sections = [
    'CalculationAuditTrail — Shows undo/redo position and last save time',
    'ValidationAlerts — Market-specific warnings (WACC < Risk-Free, Terminal Growth > WACC, etc.)',
    'Method Explainer — Educational disclaimer about fundamental vs. market sentiment',
    'ValuationStyleSelector — Conservative/Moderate/Aggressive with side-by-side DCF comparison',
    'MarketVsFundamental — Market price vs intrinsic value with key statistics',
    'ScenarioAnalysis — Bull/Bear/Base with sensitivity matrix and Monte Carlo',
    'ValuationSummaryCards — DCF, Comps, and Blended value cards',
    'BlendedWeightSlider — Adjustable DCF/Comps weighting with live preview',
    'ComparableBreakdown — Individual method implied prices (P/E, EV/EBITDA, P/S, P/B)',
    'DCFProjectionsTable — Year-by-year revenue, EBITDA, FCF, discount factor, PV',
    'KeyMetricsGrid — 12 financial ratios in card format',
    'ReverseDCFSection — Market-implied growth rate analysis',
    'QualityScorecard — 4-category quality assessment with letter grade',
    'AIReport — AI-generated investment thesis and recommendation',
]
for item in val_sections:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Charts Tab', level=3)
doc.add_paragraph(
    'Revenue/FCF/EBITDA projection chart, Valuation comparison bar chart, '
    'Football field chart (range visualization for each valuation method), '
    'and interactive Recharts components with currency-aware formatting.'
)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 13. FORMATTERS
# ════════════════════════════════════════════════════════════════
doc.add_heading('13. Formatting & Utility Functions', level=1)
doc.add_paragraph('File: src/utils/formatters.ts — Currency-aware formatting for all display values:')
add_table(
    ['Function', 'Example (USD)', 'Example (EGP)'],
    [
        ['formatCurrency()', '$1,500,000', '1,500,000 EGP'],
        ['formatCurrencyShort()', '$1.50B', '1.50B EGP'],
        ['formatPercent()', '7.50%', '7.50%'],
        ['formatMultiple()', '15.5x', '15.5x'],
        ['formatMillions()', '$1,500.00M', '1,500.00M EGP'],
        ['formatPrice()', '$45.50', '45.50 EGP'],
        ['formatShares()', '500.00M', '500.00M'],
    ]
)

doc.add_heading('14. Web Worker', level=1)
doc.add_paragraph(
    'File: src/workers/monteCarlo.worker.ts — Dedicated thread for Monte Carlo simulation. '
    'Receives financialData and assumptions via postMessage, runs 5,000 simulations, and returns '
    'statistical results. Prevents UI blocking during computation-heavy operations.'
)

doc.add_heading('15. Configuration & Build', level=1)
config_items = [
    'Vite config: React plugin + Tailwind CSS plugin + single-file plugin for deployment.',
    'Dev server: npm run dev (Vite HMR).',
    'Production build: npm run build (outputs to dist/).',
    'Environment: .env.example defines VITE_FMP_API_KEY for Financial Modeling Prep API access.',
    'TypeScript: Strict mode, ES2020 target, React JSX transform.',
]
for item in config_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
# 16. FUTURE DEVELOPMENT
# ════════════════════════════════════════════════════════════════
doc.add_heading('16. Known Issues & Future Development Areas', level=1)

doc.add_heading('Current Known Issues', level=2)
issues = [
    'Legacy valuation.ts has a duplicate WACC implementation — preferably migrate to single source in valuationEngine.ts.',
    'Monte Carlo worker duplicates logic from advancedAnalysis.ts — consider shared calculation module.',
    'Sensitivity matrix in valuationEngine.ts is incomplete (remaining lines 800-813 handle equity calc).',
    'localStorage auto-save validates JSON parsing but does not restore on mount (code commented as "could be added").',
]
for item in issues:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Suggested Future Enhancements', level=2)
future = [
    'Multi-year historical data comparison (currently only latest year)',
    'Custom peer group definition with saved presets',
    'Automated regression for beta calculation from price history',
    'Dividend Discount Model (DDM) as additional valuation method',
    'SOTP (Sum of the Parts) valuation for conglomerates',
    'Real-time price updates via WebSocket',
    'User authentication with cloud-based valuation storage',
    'Waterfall chart for EV-to-equity bridge visualization',
    'Automated PDF/Excel scheduling and email delivery',
    'Additional emerging markets (Saudi KSA, UAE, Nigeria)',
    'Unit and integration test suite',
    'Sensitivity analysis for more parameters (margin, CapEx, etc.)',
    'Mobile-responsive layout optimizations',
]
for item in future:
    doc.add_paragraph(item, style='List Bullet')

# ── Save ──
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'WOLF_Engine_Complete_Overview.docx')
doc.save(output_path)
print(f'Document saved to: {output_path}')
