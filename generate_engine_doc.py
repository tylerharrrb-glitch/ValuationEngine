"""
WOLF Valuation Engine — Complete Overview Document Generator
Generates a comprehensive Word document covering the entire engine.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import os

doc = Document()

# ── Styles ──────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_code(doc, text, lang="typescript"):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Light Shading Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for p in t.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri+1].cells[ci].text = str(val)
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(f"📌 {text}")
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x99)

# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("🐺 WOLF VALUATION ENGINE")
r.font.size = Pt(28)
r.bold = True
r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Complete Technical Overview")
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CFA-Grade 3-Statement Financial Model & DCF Engine\nOptimized for the Egyptian Market (EGX)")
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Version 2.0 — March 2026\nDeveloper Reference Document")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    "1. Executive Summary & Architecture",
    "2. Technology Stack & Project Structure",
    "3. Type System & Data Models (financial.ts)",
    "4. Market Defaults & Egyptian Tax Regimes (marketDefaults.ts)",
    "5. WACC Calculation — Dual CAPM Methodology",
    "6. FCFF Three-Way Cross-Verification",
    "7. DCF Projections & Terminal Value",
    "8. EV-to-Equity Bridge & Intrinsic Value",
    "9. Dividend Discount Models (DDM)",
    "10. Comparable Company Valuation",
    "11. Blended Valuation",
    "12. Financial Ratios (35+ Metrics)",
    "13. Sensitivity Analysis (5×5 Matrix)",
    "14. Scenario Analysis (Bear/Base/Bull)",
    "15. Reverse DCF",
    "16. Monte Carlo Simulation",
    "17. Sector Benchmarking & Quality Scorecard",
    "18. Confidence Score Module",
    "19. EAS/IFRS Compliance Modules",
    "20. Input Validation Framework",
    "21. API Integration (FMP Stable API)",
    "22. State Management & Hooks",
    "23. Export Capabilities (Excel, PDF, JSON)",
    "24. UI Components & Architecture",
    "25. Egyptian Regulatory Compliance",
    "26. File Reference Index",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1. Executive Summary & Architecture', level=1)
doc.add_paragraph(
    "The WOLF Valuation Engine is a CFA-grade, browser-based financial modeling and equity valuation platform "
    "purpose-built for the Egyptian capital market (EGX). It implements a full 3-statement model engine with "
    "circular reference resolution, multiple valuation methodologies (DCF, DDM, Comparable Companies), "
    "advanced analytics (Monte Carlo simulation, Reverse DCF, Sector Benchmarking), and Egyptian regulatory "
    "compliance (Law 159/1981 profit waterfall, Law 91/2005 tax regimes, EAS/IFRS standards)."
)
doc.add_paragraph(
    "The engine is designed as a single-page React/TypeScript application with Vite as the build tool. "
    "All financial calculations run client-side with no backend dependency. The architecture follows a strict "
    "separation between data types, calculation engines, UI components, and export formatters."
)

doc.add_heading('Architecture Overview', level=2)
add_bullet(doc, " Core calculation engine — pure functions, no side effects", "Calculation Layer:")
add_bullet(doc, " React hooks for state, memoized derived values", "State Layer:")
add_bullet(doc, " Lazy-loaded tabs (Input, Valuation, Charts) with reusable components", "UI Layer:")
add_bullet(doc, " Excel (.xlsx), PDF, JSON export with full formula fidelity", "Export Layer:")
add_bullet(doc, " FMP Stable API (2026 format) for live stock data", "Data Layer:")
add_bullet(doc, " Web Worker for Monte Carlo (5,000 simulations off main thread)", "Worker Layer:")

# ═══════════════════════════════════════════════════════════════
# 2. TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2. Technology Stack & Project Structure', level=1)
doc.add_heading('Stack', level=2)
add_table(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Framework', 'React 18 + TypeScript', 'UI rendering & type safety'],
        ['Build Tool', 'Vite', 'Fast HMR & optimized builds'],
        ['Styling', 'Tailwind CSS', 'Utility-first responsive design'],
        ['Charts', 'Recharts', 'Revenue/FCF/Sensitivity visualizations'],
        ['Excel Export', 'ExcelJS', 'Multi-sheet .xlsx with formulas'],
        ['PDF Export', 'jsPDF + html2canvas', 'Professional PDF reports'],
        ['API', 'Financial Modeling Prep (FMP)', 'Live market data ingestion'],
        ['State', 'React useState + custom hooks', 'Client-side state with history'],
        ['Worker', 'Web Worker API', 'Monte Carlo off main thread'],
        ['Hosting', 'Cloudflare Pages', 'Static site deployment'],
    ]
)
doc.add_paragraph()

doc.add_heading('Project Structure', level=2)
add_code(doc, """ValuationEngine/
├── src/
│   ├── App.tsx                      # Main orchestrator (177 lines)
│   ├── main.tsx                     # React entry point
│   ├── types/
│   │   └── financial.ts             # All TypeScript interfaces (465 lines)
│   ├── constants/
│   │   ├── initialData.ts           # Default sample data & assumptions
│   │   ├── marketDefaults.ts        # USA/Egypt market configs & tax categories
│   │   └── valuationStyles.ts       # Conservative/Moderate/Aggressive presets
│   ├── utils/
│   │   ├── valuationEngine.ts       # CORE ENGINE (1,098 lines)
│   │   ├── advancedAnalysis.ts      # Reverse DCF, Monte Carlo, Benchmarks, Quality (679 lines)
│   │   ├── confidenceScore.ts       # 3-pillar model confidence (186 lines)
│   │   ├── valuation.ts             # Convenience wrappers (207 lines)
│   │   ├── easModules.ts            # EAS 48/31/12/23 compliance (232 lines)
│   │   ├── inputValidation.ts       # Hard blocks, warnings, edge cases (353 lines)
│   │   ├── excelExport.ts           # Multi-sheet Excel export (52K)
│   │   ├── excelExportPro.ts        # Pro Excel with formulas (49K)
│   │   ├── pdfExport.ts             # Professional PDF reports (80K)
│   │   ├── jsonExport.ts            # Structured JSON export
│   │   ├── formatters.ts            # USD/EGP currency formatting (203 lines)
│   │   ├── industryMapping.ts       # Ticker → industry mapping
│   │   └── calculations/
│   │       ├── dcf.ts               # DCF projection helpers
│   │       ├── comparables.ts       # Comparable company helpers
│   │       ├── metrics.ts           # Key metrics & recommendation
│   │       └── scenarios.ts         # Scenario case calculations
│   ├── hooks/
│   │   ├── useFinancialData.ts      # State management with undo/redo & auto-save
│   │   ├── useValuationCalculations.ts  # Memoized calculation pipeline
│   │   ├── useHistory.ts            # Undo/redo history stack
│   │   ├── useTheme.ts             # Dark/light mode
│   │   └── useKeyboardShortcuts.ts  # Ctrl+Z, Ctrl+Y, Ctrl+S
│   ├── services/
│   │   ├── stockAPI.ts              # FMP Stable API client (729 lines)
│   │   └── yahooFinanceAPI.ts       # Yahoo Finance fallback
│   ├── workers/
│   │   └── monteCarlo.worker.ts     # Web Worker for Monte Carlo (143 lines)
│   └── components/
│       ├── layout/                  # Header, Footer, TabNavigation, CompanyHeader
│       ├── input/                   # InputTab (financial data entry forms)
│       ├── valuation/               # ValuationTab + sections (DCF, Comps, DDM...)
│       ├── charts/                  # ChartsTab (Recharts visualizations)
│       └── shared/                  # ErrorBoundary, LoadingFallback
│       ├── StockSearch.tsx          # Ticker search with API integration
│       ├── AssumptionsPanel.tsx      # WACC, CAPM, projection drivers
│       ├── DCFResults.tsx           # DCF output display
│       ├── ComparableCompanies.tsx  # Peer comparison table
│       ├── FootballFieldChart.tsx   # Multi-methodology range chart
│       ├── KeyMetrics.tsx           # Financial ratios dashboard
│       ├── ValuationSummary.tsx     # Verdict & recommendation
│       ├── ScenarioToggle.tsx       # Bear/Base/Bull selector
│       ├── AIReport.tsx             # AI-powered analysis report
│       ├── UserAuth.tsx             # User authentication
│       └── ...                      # 17 total standalone components
""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. TYPE SYSTEM
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3. Type System & Data Models', level=1)
doc.add_paragraph("File: src/types/financial.ts (465 lines)")
doc.add_paragraph(
    "The entire engine is strictly typed in TypeScript. Every financial input, assumption, intermediate "
    "calculation, and output has a corresponding interface. This ensures compile-time safety across all "
    "calculation paths, export formatters, and UI components."
)

doc.add_heading('Core Interfaces', level=2)
add_table(doc,
    ['Interface', 'Fields', 'Description'],
    [
        ['IncomeStatement', '10', 'Revenue, COGS, Gross Profit, OpEx, EBIT, Interest, Tax, NI, D&A'],
        ['BalanceSheet', '18+', 'Current/Non-current assets, liabilities, equity, optional EAS fields'],
        ['CashFlowStatement', '5+', 'Operating CF, CapEx, FCF, Dividends Paid, Net Change'],
        ['FinancialData', '10+', 'Company name, ticker, shares, price, IS/BS/CF, historical data'],
        ['ValuationAssumptions', '30+', 'All DCF/CAPM/DDM/Scenario parameters'],
        ['DCFProjection', '11', 'Year, Revenue, EBITDA, D&A, EBIT, NOPAT, CapEx, ΔWC, FCFF, DF, PV'],
        ['WACCResult', '10', 'Ke, Kd(1-t), weights, WACC, Market Cap, Debt, CAPM method'],
        ['FCFFVerification', '5', 'Method 1/2/3 values, allMatch boolean, tolerance'],
        ['DCFResult', '11', 'Sum PV, TV, PV(TV), TV%, EV, Net Debt, Equity, Price, Upside, Verdict'],
        ['DDMResult', '5', 'Gordon Growth, Two-Stage, H-Model values, applicable flag'],
        ['SensitivityMatrix', '3', 'WACC axis, Growth axis, 5×5 cell grid'],
        ['ScenarioAnalysis', '4', 'Bear/Base/Bull cases + weighted value'],
        ['FinancialRatios', '35+', 'Profitability, Leverage, Liquidity, Efficiency, Valuation multiples'],
        ['ValidationAlert', '4', 'type (error/warning/info), message, field, blocking flag'],
        ['ValuationJSON', 'nested', 'Full export schema with metadata, inputs, all calculations'],
    ]
)
doc.add_paragraph()

doc.add_heading('Enumerated Types', level=2)
add_table(doc,
    ['Type', 'Values', 'Purpose'],
    [
        ['CAPMMethod', "'A' | 'B'", 'Local Currency CAPM vs USD Build-Up'],
        ['EgyptTaxCategory', "5 categories", 'standard, oil_gas, suez_canal, free_zone, custom'],
        ['BetaType', "'raw' | 'adjusted' | 'relevered'", 'Beta calculation methodology'],
        ['DiscountingConvention', "'end_of_year' | 'mid_year'", 'Cash flow timing assumption'],
        ['TerminalValueMethod', "'gordon_growth' | 'exit_multiple'", 'Terminal value approach'],
        ['MarketRegion', "'USA' | 'Egypt'", 'Market configuration selector'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. MARKET DEFAULTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4. Market Defaults & Egyptian Tax Regimes', level=1)
doc.add_paragraph("File: src/constants/marketDefaults.ts (96 lines)")

doc.add_heading('Dual-Market Configuration', level=2)
add_table(doc,
    ['Parameter', 'USA', 'Egypt', 'Source'],
    [
        ['Risk-Free Rate', '4.5%', '22.0%', '10Y Treasury / 10Y EG Bond'],
        ['Market Risk Premium', '5.5%', '5.5%', 'Damodaran Mature ERP'],
        ['Country Risk Premium', '0%', '7.5% (Method B only)', 'Damodaran Egypt CRP'],
        ['Terminal Growth', '2.5%', '8.0%', 'Nominal GDP proxy'],
        ['Max Terminal Growth', '4.0%', '12.0%', 'Sustainable cap'],
        ['Default Tax Rate', '21.0%', '22.5%', 'Statutory corporate tax'],
        ['Currency', 'USD ($)', 'EGP', '—'],
        ['CBE Benchmark Rate', 'N/A', '27.25%', 'CBE overnight lending corridor'],
        ['Dividend WHT', 'N/A', '10%', 'Article 46 bis'],
        ['Capital Gains Tax', 'N/A', '10%', 'On listed securities'],
    ]
)
doc.add_paragraph()

doc.add_heading('Egyptian Tax Categories (Law 91/2005)', level=2)
add_table(doc,
    ['Category', 'Rate', 'Applies To'],
    [
        ['Standard (DEFAULT)', '22.5%', 'Most Egyptian companies'],
        ['Industrial Zone', '0% (5 years)', 'Industrial free zone — reverts to 22.5%'],
        ['Oil & Gas', '40.55%', 'Petroleum sector companies'],
        ['Suez Canal / EGPC / CBE', '40.0%', 'Specific state entities'],
        ['Free Zone (Export)', '0%', 'Free economic zone exports'],
        ['Custom', 'User-defined', 'Any special case'],
    ]
)
doc.add_paragraph()

doc.add_heading('Egyptian Industry Multiples (EGX Defaults)', level=2)
add_table(doc,
    ['Sector', 'P/E', 'EV/EBITDA', 'P/S', 'P/B'],
    [
        ['Banking', '5.5x', 'N/A', '2.0x', '1.0x'],
        ['Real Estate', '8.0x', '7.0x', '1.5x', '0.8x'],
        ['Telecom', '12.0x', '6.0x', '1.8x', '2.0x'],
        ['Consumer/FMCG', '15.0x', '8.0x', '1.0x', '2.5x'],
        ['Industrial', '7.0x', '5.0x', '0.8x', '1.2x'],
        ['Healthcare', '18.0x', '10.0x', '2.0x', '3.0x'],
        ['Energy', '6.0x', '4.5x', '0.6x', '1.0x'],
        ['EGX Market Average', '7.0x', '5.0x', '1.2x', '1.5x'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. WACC CALCULATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5. WACC Calculation — Dual CAPM Methodology', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — calculateWACC(), calculateCostOfEquity()")
doc.add_paragraph(
    "The WACC is always calculated from its components — never hard-coded. Capital structure weights "
    "use MARKET CAPITALIZATION (not book equity) per CFA best practices."
)

doc.add_heading('Master Formula', level=2)
add_code(doc, "WACC = (E/V) × Ke + (D/V) × Kd × (1 − t)\nwhere V = E + D = Market Cap + Total Debt")

doc.add_heading('Method A — Local Currency CAPM (Default for EGP)', level=2)
add_code(doc, "Ke = Rf(Egypt) + β × Mature_Market_ERP\n   = 22.0% + β × 5.5%\n\n⚠️ NO Country Risk Premium — it is already embedded in the Egyptian Rf")
doc.add_paragraph("This is the default method. The Egyptian risk-free rate (10-year government bond) already incorporates sovereign risk, so adding CRP would double-count.")

doc.add_heading('Method B — USD Build-Up (Alternative)', level=2)
add_code(doc, """Ke(USD) = Rf(US) + β × Mature_ERP + CRP
        = 4.5% + β × 5.5% + 7.5%

Ke(EGP) = [(1 + Ke_USD) × (1 + Egypt_Inflation) / (1 + US_Inflation)] − 1
        (Fisher equation for currency conversion)""")

doc.add_heading('Beta Adjustments', level=2)
add_table(doc,
    ['Type', 'Formula', 'Use Case'],
    [
        ['Raw', 'β (as observed)', 'Default — regression beta vs EGX 30 or S&P 500'],
        ['Adjusted (Bloomberg)', '⅔ × Raw β + ⅓ × 1.0', 'Mean-reversion adjustment'],
        ['Relevered', 'βL = βU × [1 + (1−t) × (D/E)]', 'Adjusting unlevered sector beta for company leverage'],
    ]
)

doc.add_heading('After-Tax Cost of Debt', level=2)
add_code(doc, "Kd(after-tax) = Kd(pre-tax) × (1 − t)\nDefault Egypt: 20.0% × (1 − 0.225) = 15.5%")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. FCFF
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6. FCFF Three-Way Cross-Verification', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — calculateFCFFVerification()")
doc.add_paragraph("All three methods MUST produce identical results within a tolerance of ±0.001. This is a hard integrity check.")

doc.add_heading('Three Methods', level=2)
add_code(doc, """Method 1 — NOPAT Route (Primary):
  FCFF = EBIT × (1−t) + D&A − CapEx − ΔWC

Method 2 — EBITDA Route:
  FCFF = EBITDA × (1−t) + D&A × t − CapEx − ΔWC

Method 3 — Net Income Route:
  FCFF = Net Income + Interest × (1−t) + D&A − CapEx − ΔWC""")

add_note(doc, "FCFF must NEVER subtract Interest Expense directly. ΔWC increase = cash OUTFLOW (subtracted).")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. DCF PROJECTIONS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7. DCF Projections & Terminal Value', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — calculateDCFProjections(), calculateDCFValue()")

doc.add_heading('Year-by-Year Projection Chain', level=2)
add_code(doc, """For each year t = 1 to N (where N ∈ {3, 5, 7, 10}):
  Revenue(t)  = Revenue(t-1) × (1 + Revenue_Growth)
  EBITDA(t)   = Revenue(t) × EBITDA_Margin
  D&A(t)      = Revenue(t) × DA_Percent
  EBIT(t)     = EBITDA(t) − D&A(t)
  NOPAT(t)    = EBIT(t) × (1 − Tax_Rate)
  CapEx(t)    = Revenue(t) × CapEx_Percent
  ΔWC(t)      = [Revenue(t) − Revenue(t-1)] × DeltaWC_Percent
  FCFF(t)     = NOPAT(t) + D&A(t) − CapEx(t) − ΔWC(t)""")

doc.add_heading('Discounting Convention', level=2)
add_table(doc,
    ['Convention', 'Discount Factor', 'When to Use'],
    [
        ['End-of-Year (default)', '(1 + WACC)^t', 'Standard assumption — cash flows arrive at year end'],
        ['Mid-Year', '(1 + WACC)^(t − 0.5)', 'Cash flows arrive evenly throughout the year (more accurate)'],
    ]
)

doc.add_heading('Terminal Value', level=2)
add_code(doc, """Gordon Growth (default):
  TV = FCFF_N × (1 + g) / (WACC − g)
  ⚠️ Hard requirement: g < WACC (otherwise model returns 0)

Exit Multiple:
  TV = EBITDA_N × Exit_Multiple

PV(TV) = TV / (1 + WACC)^N""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. EV-to-Equity Bridge
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8. EV-to-Equity Bridge & Intrinsic Value', level=1)
add_code(doc, """Enterprise Value = Σ PV(FCFF) + PV(Terminal Value)

Equity Value = Enterprise Value − Total Debt + Cash
             = EV − (Short-term Debt + Long-term Debt) + Cash & Equivalents

⚠️ NO MAX(0) floor — Equity Value CAN be negative for distressed companies

Intrinsic Value per Share = Equity Value / Shares Outstanding""")

doc.add_heading('Verdict Logic', level=2)
add_table(doc,
    ['Verdict', 'Condition', 'Signal'],
    [
        ['UNDERVALUED', 'Intrinsic > Market Price × 1.10', '>10% upside'],
        ['FAIRLY VALUED', '0.90 × Market ≤ Intrinsic ≤ 1.10 × Market', 'Within ±10% band'],
        ['OVERVALUED', 'Intrinsic < Market Price × 0.90', '>10% downside'],
    ]
)

# ═══════════════════════════════════════════════════════════════
# 9. DDM
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9. Dividend Discount Models (DDM)', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — calculateDDM()")
doc.add_paragraph("All DDM models discount at Ke (cost of equity), NOT WACC. DDM returns N/A if company pays no dividends or has negative earnings.")

doc.add_heading('Three DDM Variants', level=2)
add_code(doc, """9.1 Gordon Growth (Single-Stage):
  P = D₀ × (1 + g) / (Ke − g) = D₁ / (Ke − g)

9.2 Two-Stage DDM:
  P = Σ[D₀(1+g₁)^t / (1+Ke)^t] for t=1..n
    + [Dₙ(1+g₂)/(Ke−g₂)] / (1+Ke)^n

9.3 H-Model:
  P = D₀(1+gL)/(Ke−gL) + D₀×H×(gS−gL)/(Ke−gL)
  where H = n/2 (half-life of high-growth period)""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 10. COMPARABLE COMPANY
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10. Comparable Company Valuation', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — calculateComparableValue()")

doc.add_heading('Four Multiples Used', level=2)
add_table(doc,
    ['Multiple', 'Implied Price Formula', 'Notes'],
    [
        ['P/E', 'EPS × Median Peer P/E', 'Only positive EPS'],
        ['EV/EBITDA', '(EBITDA × Med. EV/EBITDA − Debt + Cash) / Shares', 'Enterprise to equity bridge — no MAX(0) floor'],
        ['P/S', 'Revenue/Share × Median P/S', 'Useful for pre-profit companies'],
        ['P/B', 'Book Value/Share × Median P/B', 'Asset-intensive sectors'],
    ]
)
doc.add_paragraph("Peer companies can be sourced via FMP API auto-fetch or manually entered. Median (not mean) is used for robustness against outliers. Both US and Egyptian peer groups are pre-configured.")

# ═══════════════════════════════════════════════════════════════
# 11. BLENDED VALUATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11. Blended Valuation', level=1)
add_code(doc, """Blended Value = (DCF Value × DCF_Weight) + (Comparable Value × Comps_Weight)
Default weights: DCF 60%, Comps 40% (adjustable via slider)

If no valid comparables exist, DCF weight = 100%.

Upside = (Blended − Current Price) / Current Price × 100%""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 12. FINANCIAL RATIOS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12. Financial Ratios (35+ Metrics)', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — calculateFinancialRatios()")

add_table(doc,
    ['Category', 'Metrics', 'Key Formulas'],
    [
        ['Profitability', 'Gross Margin, EBITDA Margin, EBIT Margin, Net Margin, ROE, ROA, ROIC', 'ROIC = NOPAT / Invested Capital'],
        ['Leverage', 'D/E, Net Debt/EBITDA, Interest Coverage', 'IC = EBIT / Interest Expense'],
        ['Liquidity', 'Current Ratio, Quick Ratio', 'Quick = (Cash + AR) / Current Liabilities'],
        ['Efficiency', 'DSO, DIO, DPO, Cash Conversion Cycle', 'CCC = DSO + DIO − DPO'],
        ['Value Creation', 'EVA, ROIC-WACC Spread', 'EVA = NOPAT − (IC × WACC)'],
        ['Quality', 'Altman Z-Score, Piotroski F-Score (TODO)', 'Z = 1.2×WC/TA + 1.4×RE/TA + 3.3×EBIT/TA + 0.6×MCap/TL + Rev/TA'],
        ['Valuation', 'P/E (trailing & forward), EV/EBITDA, EV/EBIT, EV/Revenue, P/B, P/S, P/CF, PEG, Dividend Yield, Earnings Yield, FCF Yield', 'Forward P/E uses 15% growth proxy'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 13-16. ADVANCED ANALYTICS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13. Sensitivity Analysis (5×5 Matrix)', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — generateSensitivityMatrix()")
add_code(doc, """WACC Axis: [Base−4%, Base−2%, Base, Base+2%, Base+4%]
Growth Axis: [Base−3%, Base−1.5%, Base, Base+1.5%, Base+3%]
  Growth floored at 4% (Egypt minimum), capped below WACC−1%

Each cell recalculates full DCF with those parameters.
Color coding: Green (>10% upside), Yellow (±10%), Red (>10% downside)
Base case cell is highlighted.""")

doc.add_heading('14. Scenario Analysis (Bear/Base/Bull)', level=1)
doc.add_paragraph("File: src/utils/valuationEngine.ts — calculateScenarioAnalysis()")
add_table(doc,
    ['Parameter', 'Bear Delta', 'Base', 'Bull Delta'],
    [
        ['Revenue Growth', '−5%', '0', '+5%'],
        ['EBITDA Margin', '−3%', '0', '+3%'],
        ['Terminal Growth', '−1%', '0', '+1%'],
        ['WACC', '+1%', '0', '−1%'],
        ['Default Probability', '25%', '50%', '25%'],
    ]
)
add_code(doc, "Weighted Value = P(Bear)×V(Bear) + P(Base)×V(Base) + P(Bull)×V(Bull)")

doc.add_heading('15. Reverse DCF', level=1)
doc.add_paragraph("File: src/utils/advancedAnalysis.ts — calculateReverseDCF()")
doc.add_paragraph("Uses binary search (100 iterations) to solve for the implied revenue growth rate the market is pricing in. Compares to your base case growth to classify market expectations as 'aggressive', 'reasonable', or 'conservative'.")

doc.add_heading('16. Monte Carlo Simulation', level=1)
doc.add_paragraph("Files: src/utils/advancedAnalysis.ts, src/workers/monteCarlo.worker.ts")
add_code(doc, """Runs 5,000 simulations with Gaussian random variation:
  Revenue Growth: ±4% std dev
  WACC: ±1.5% std dev
  Terminal Growth: ±0.8% std dev
  EBITDA Margin: ±1.5% std dev

Outputs: Mean/Median price, Std Dev, P5/P25/P75/P95 percentiles,
         probability above current price, 20-bucket histogram distribution.

Worker thread: Runs off main thread via Web Worker API to prevent UI blocking.""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 17-18. QUALITY & CONFIDENCE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('17. Sector Benchmarking & Quality Scorecard', level=1)
doc.add_paragraph("File: src/utils/advancedAnalysis.ts")

doc.add_heading('Sector Benchmarking', level=2)
doc.add_paragraph("Compares the company against sector median/P25/P75 benchmarks across 10 metrics: Gross Margin, Operating Margin, Net Margin, ROE, ROA, FCF Margin, D/E, Current Ratio, P/E, EV/EBITDA. Generates an overall percentile score and qualitative rating.")

doc.add_heading('Quality Scorecard (40 points total)', level=2)
add_table(doc,
    ['Pillar', 'Max Score', 'Components'],
    [
        ['Economic Moat', '10', 'Gross Margin, Operating Margin, ROE, Market Position'],
        ['Financial Health', '10', 'Current Ratio, D/E, Interest Coverage, Cash/Debt'],
        ['Growth & Profitability', '10', 'Net Margin, FCF Generation, Earnings Quality (OCF>NI)'],
        ['Capital Allocation', '10', 'CapEx/Revenue, Dividend Payout, FCF Conversion, Cash Accumulation'],
    ]
)
add_code(doc, "Grade: A+ (≥85%) → +15% premium | A (≥75%) → +10% | B+ (≥65%) → +5% | B (≥55%) → 0%\n        C (≥45%) → −5% | D (≥35%) → −10% | F (<35%) → −15%")

doc.add_heading('18. Confidence Score Module', level=1)
doc.add_paragraph("File: src/utils/confidenceScore.ts (186 lines)")
add_table(doc,
    ['Pillar', 'Max Points', 'Checks'],
    [
        ['Data Quality', '30', 'D&A>0 (5), CapEx>0 (5), Interest consistency (5), EBITDA>0 (5), Revenue>0 (5), Shares>0 (5)'],
        ['Assumption Reasonableness', '40', 'Rev growth ≤25% (8), EBITDA margin ≤50% (8), g<GDP (8), WACC 15-35% (8), FCFF reconciles (8)'],
        ['Model Robustness', '30', 'TV<80% EV (10), Value within ±50% (10), DCF/DDM agreement (5), IC>1.5x (5)'],
    ]
)
add_code(doc, "Grades: A (≥85) | B (≥70) | C (≥55) | D (≥40) | F (<40)")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 19. EAS COMPLIANCE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('19. EAS/IFRS Compliance Modules', level=1)
doc.add_paragraph("File: src/utils/easModules.ts (232 lines)")

doc.add_heading('EAS 48 (IFRS 16) — Lease Adjustments', level=2)
add_code(doc, """ROU Asset = PV of remaining lease payments at IBR
         = PMT × [(1 − (1+r)^−n) / r]
Lease Liability = ROU Asset (at inception)
Annual Depreciation = ROU / remaining term (straight-line)
Interest Expense = Lease Liability × IBR (Year 1)
EBITDA Add-back = Full annual lease payment (reclassified from operating to financing)
Post-adjustment Debt = Current Debt + Lease Liability""")

doc.add_heading('EAS 31 (IAS 1) — Normalized Earnings', level=2)
doc.add_paragraph("Add-backs (impairment, restructuring, legal settlements) and deductions (gain on asset sale, insurance proceeds) with optional tax-effect adjustment.")
add_code(doc, "Normalized NI = Reported NI + Σ(Adjustments) − Tax Effect on tax-affected items")

doc.add_heading('EAS 12 (IAS 12) — Deferred Tax in EV Bridge', level=2)
add_code(doc, "Net DTA = Deferred Tax Asset − Deferred Tax Liability\nAdjusted Equity = Base Equity + Net DTA")

doc.add_heading('EAS 23 (IAS 33) — Basic & Diluted EPS', level=2)
add_code(doc, """Basic EPS = Net Income / Basic Shares
Diluted EPS = (NI + Convertible Interest) / Diluted Shares
Anti-dilution check: if Diluted EPS > Basic EPS, use Basic EPS""")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 20. INPUT VALIDATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('20. Input Validation Framework', level=1)
doc.add_paragraph("File: src/utils/inputValidation.ts (353 lines)")

doc.add_heading('Hard Blocks (Must Fix — Calculation Blocked)', level=2)
add_table(doc,
    ['ID', 'Rule', 'Limit'],
    [
        ['HB-1', 'Revenue ≥ 0', '≥ 0'],
        ['HB-2', 'Tax Rate', '0–60%'],
        ['HB-3', 'Risk-Free Rate', '0–40%'],
        ['HB-4', 'Equity Risk Premium', '0–20%'],
        ['HB-5', 'Cost of Debt', '0–50%'],
        ['HB-6', 'Terminal Growth < WACC', 'g < WACC (hard block)'],
        ['HB-7', 'Shares Outstanding > 0', '> 0'],
        ['HB-8', 'Projection Years ∈ {3,5,7,10}', 'Fixed set'],
    ]
)

doc.add_heading('Soft Warnings (Calculation Proceeds)', level=2)
add_table(doc,
    ['ID', 'Rule'],
    [
        ['W-1', 'Negative EBITDA — distress signal'],
        ['W-2', 'Beta outside 0–2.5'],
        ['W-3', 'Terminal growth > 10%'],
        ['W-4', 'CapEx < D&A — underinvestment'],
        ['W-5', 'Interest Coverage < 1.5x'],
        ['W-6', 'Dividend payout > 100%'],
        ['W-7', 'Revenue growth > 25%'],
        ['W-8', 'Terminal Value > 85% of EV'],
        ['W-9', 'Negative book equity'],
    ]
)

doc.add_heading('Edge Cases (Auto-Handled)', level=2)
add_table(doc,
    ['ID', 'Case', 'Engine Behavior'],
    [
        ['EC-1', 'All-equity firm (no debt)', 'WACC = Ke'],
        ['EC-2', 'Net cash position', 'EV < Market Cap'],
        ['EC-3', 'Zero dividends', 'DDM returns N/A'],
        ['EC-4', 'WACC > 30%', 'Normal for Egypt — no adjustment'],
        ['EC-5', 'Exit multiple method', 'TV = EBITDA × multiple (not Gordon)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 21. API INTEGRATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('21. API Integration (FMP Stable API)', level=1)
doc.add_paragraph("File: src/services/stockAPI.ts (729 lines)")

doc.add_heading('Architecture', level=2)
doc.add_paragraph(
    "Uses Financial Modeling Prep (FMP) NEW Stable API (2026 format) with query parameters. "
    "All four endpoints are fetched in parallel via Promise.all for speed."
)
add_table(doc,
    ['Endpoint', 'Data Retrieved'],
    [
        ['profile', 'Company name, price, market cap, beta, shares outstanding, sector'],
        ['income-statement', 'Revenue, COGS, EBIT, Interest, Tax, Net Income, D&A, EBITDA'],
        ['balance-sheet-statement', 'All asset/liability/equity line items'],
        ['cash-flow-statement', 'Operating CF, CapEx, FCF, Dividends Paid, Net Change in Cash'],
    ]
)

doc.add_heading('Smart Data Handling', level=2)
add_bullet(doc, ": Shares Outstanding fallback chain: profile.sharesOutstanding → mktCap/price → 1B default", "Shares")
add_bullet(doc, ": If API reports 0 interest but debt exists, assumes 2.8% implied cost", "Interest")
add_bullet(doc, ": If sector is Financial and COGS=0, recalculates Gross Profit as Revenue − Interest Expense", "Banks")
add_bullet(doc, ": .CA suffix auto-sets market to Egypt with Egyptian defaults", "Egyptian Stocks")
add_bullet(doc, ": If unclassified non-current assets > 40% of total, reclassifies as Long-term Investments", "LT Investments")
add_bullet(doc, ": FMP reports outflows as negative — engine stores dividends/capex as positive", "Dividends")

doc.add_heading('Peer Company System', level=2)
doc.add_paragraph("Pre-configured peer groups for both US (AAPL, MSFT, GOOGL, etc.) and Egyptian companies (COMI.CA, ETEL.CA, EFID.CA, TMGH.CA). Auto-suggests peers based on ticker lookup. Fetches ratios-ttm endpoint for live multiples.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 22. STATE MANAGEMENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('22. State Management & Hooks', level=1)

doc.add_heading('useFinancialData (215 lines)', level=2)
add_bullet(doc, "Manages financialData, assumptions, and comparables state")
add_bullet(doc, "Every update saves to undo/redo history stack")
add_bullet(doc, "Auto-saves to localStorage with 1-second debounce")
add_bullet(doc, "handleStockDataFetched: auto-detects Egyptian stocks (.CA suffix), sets market defaults, calculates tax rate and WACC from API data")

doc.add_heading('useValuationCalculations (209 lines)', level=2)
add_bullet(doc, "All calculations wrapped in useMemo — only recompute when inputs change")
add_bullet(doc, "Applies both Scenario (Bear/Base/Bull) AND Valuation Style (Conservative/Moderate/Aggressive) multipliers")
add_bullet(doc, "Generates Football Field chart data across DCF, P/E, EV/EBITDA, DDM, P/B, and Blended methods")
add_bullet(doc, "Auto-detects financial sector companies and hides EV/EBITDA bar from Football Field chart")

doc.add_heading('useHistory, useTheme, useKeyboardShortcuts', level=2)
add_bullet(doc, ": 50-state undo/redo stack with Ctrl+Z/Ctrl+Y", "useHistory")
add_bullet(doc, ": Persisted dark/light mode toggle", "useTheme")
add_bullet(doc, ": Ctrl+Z (undo), Ctrl+Y (redo), Ctrl+S (export PDF)", "useKeyboardShortcuts")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 23. EXPORT CAPABILITIES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('23. Export Capabilities', level=1)

doc.add_heading('Excel Export (52K + 49K lines)', level=2)
doc.add_paragraph("Two versions: Standard (data-only) and Pro (with live Excel formulas).")
add_bullet(doc, "Multi-sheet workbook: Dashboard, Income Statement, Balance Sheet, Cash Flow, DCF, Ratios, Scenarios, Sensitivity")
add_bullet(doc, "Pro version uses _Calc sheets pattern for scenario switching")
add_bullet(doc, "Egyptian profit waterfall formulas (Legal Reserve, EPD, Dividends, Retained Earnings) replicated in Excel")
add_bullet(doc, "Color-coded formatting with conditional color scales")

doc.add_heading('PDF Export (80K lines)', level=2)
add_bullet(doc, "Professional multi-page report with cover page, table of contents, and charts")
add_bullet(doc, "Includes all valuation results, sensitivity matrix, scenario analysis, and financial ratios")

doc.add_heading('JSON Export', level=2)
add_bullet(doc, "Structured ValuationJSON schema with metadata (engine version, date, currency, CAPM method)")
add_bullet(doc, "Full export of all inputs, calculated results, projections, and verification data")
add_bullet(doc, "Designed for data portability and external system integration")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 24. UI COMPONENTS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('24. UI Components & Architecture', level=1)
doc.add_paragraph("The app uses 3 lazy-loaded tab views with React.lazy() and Suspense for code splitting:")

doc.add_heading('Tab Architecture', level=2)
add_table(doc,
    ['Tab', 'Component', 'Contents'],
    [
        ['Input', 'InputTab', 'Stock search, financial data forms, assumptions panel, comparable companies'],
        ['Valuation', 'ValuationTab', 'DCF results, DDM, Comps, Blended, Sensitivity matrix, Scenarios, Football Field, Quality Scorecard, Confidence Score'],
        ['Charts', 'ChartsTab', 'Revenue/FCF projection charts, Valuation comparison bar chart, Monte Carlo histogram, Sector benchmark radar'],
    ]
)

doc.add_heading('Key Components', level=2)
add_table(doc,
    ['Component', 'Purpose'],
    [
        ['StockSearch.tsx', 'Ticker input with FMP API integration, auto-populate all financials'],
        ['AssumptionsPanel.tsx', 'CAPM method toggle, beta type, tax category, projection drivers'],
        ['DCFResults.tsx', 'Full DCF output with FCFF verification, TV breakdown, EV bridge'],
        ['ComparableCompanies.tsx', 'Peer table with auto-fetch, median calculation'],
        ['FootballFieldChart.tsx', 'Multi-method valuation range visualization'],
        ['KeyMetrics.tsx', 'Financial ratios dashboard with color-coded ratings'],
        ['ValuationSummary.tsx', 'Final verdict, blended value, upside/downside'],
        ['ScenarioToggle.tsx', 'Bear/Base/Bull with multiplier presets'],
        ['AIReport.tsx', 'AI-powered narrative analysis report'],
        ['Tooltip.tsx', 'Educational tooltips with formula explanations'],
        ['WolfLogo.tsx', 'Animated WOLF brand logo'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 25. EGYPTIAN REGULATORY COMPLIANCE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('25. Egyptian Regulatory Compliance', level=1)

doc.add_heading('Law 159/1981 — Profit Distribution Waterfall', level=2)
add_code(doc, """1. Net Income (After Corporate Tax)         ← Starting point
2. Legal Reserve = 5% of NI                 ← Cap: 50% of Paid-Up Capital
3. Distributable Profit = NI − Legal Reserve
4. Employee Profit Distribution (EPD) = 10% of Distributable Profit
                                        ← Cap: Total Annual Payroll
5. NI After EPD = Distributable Profit − EPD
6. Dividends = NI After EPD × Payout Ratio  ← Cap: Distributable Profit
7. Addition to Retained Earnings = NI After EPD − Gross Dividends""")

doc.add_heading('Law 91/2005 & Law 6/2025 — Tax Regimes', level=2)
add_table(doc,
    ['Regime', 'Rate', 'Notes'],
    [
        ['Standard Corporate', '22.5%', 'Default for most companies'],
        ['Oil Exploration', '40.55%', 'Petroleum sector'],
        ['Strategic Entities / CBE / Suez Canal', '40.0%', 'State entities'],
        ['Investment Law 72 Zone A', '50% deduction', 'Of investment cost from taxable income'],
        ['Investment Law 72 Zone B', '30% deduction', 'Of investment cost from taxable income'],
    ]
)

doc.add_heading('SME Turnover Tax (Law 6/2025)', level=2)
add_table(doc,
    ['Revenue Bracket (EGP)', 'Rate'],
    [
        ['Up to 1,000,000', '0.40%'],
        ['1,000,001 – 2,000,000', '0.75%'],
        ['2,000,001 – 3,000,000', '1.00%'],
        ['3,000,001 – 20,000,000', '1.50%'],
    ]
)
doc.add_paragraph("Note: SME tax applies to REVENUE, not Earnings Before Tax.")

doc.add_heading('Other Specifics', level=2)
add_bullet(doc, ": 5% for EGX listed, 10% for unlisted", "Dividend Withholding Tax")
add_bullet(doc, ": Max 5 years (Law 91/2005). Taxable income clamped at MAX(0, EBT − LossUtilized)", "Loss Carryforward")
add_bullet(doc, ": 10% on listed securities", "Capital Gains Tax")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 26. FILE REFERENCE INDEX
# ═══════════════════════════════════════════════════════════════
doc.add_heading('26. File Reference Index', level=1)
add_table(doc,
    ['File', 'Lines', 'Purpose'],
    [
        ['src/types/financial.ts', '465', 'All TypeScript interfaces and type definitions'],
        ['src/utils/valuationEngine.ts', '1,098', 'CORE: WACC, CAPM, FCFF, DCF, DDM, Comps, Blended, Ratios, Sensitivity, Scenarios, Reverse DCF, Validation'],
        ['src/utils/advancedAnalysis.ts', '679', 'Reverse DCF (binary search), Monte Carlo, Sector Benchmarks, Quality Scorecard'],
        ['src/utils/confidenceScore.ts', '186', '3-pillar model confidence scoring (0-100)'],
        ['src/utils/valuation.ts', '207', 'Convenience wrappers around valuationEngine.ts'],
        ['src/utils/easModules.ts', '232', 'EAS 48 (IFRS 16), EAS 31 (IAS 1), EAS 12 (IAS 12), EAS 23 (IAS 33)'],
        ['src/utils/inputValidation.ts', '353', 'Hard blocks (HB-1 to HB-8), Warnings (W-1 to W-9), Edge cases (EC-1 to EC-5)'],
        ['src/utils/excelExport.ts', '~1,500', 'Multi-sheet Excel export with data formatting'],
        ['src/utils/excelExportPro.ts', '~1,400', 'Pro Excel with live formulas and _Calc sheet pattern'],
        ['src/utils/pdfExport.ts', '~2,200', 'Professional multi-page PDF report generation'],
        ['src/utils/jsonExport.ts', '~160', 'Structured JSON export with full metadata'],
        ['src/utils/formatters.ts', '203', 'USD/EGP currency formatting, number formatting, chart labels'],
        ['src/utils/industryMapping.ts', '~50', 'Ticker to industry/sector mapping'],
        ['src/constants/marketDefaults.ts', '96', 'USA/Egypt market configs, tax categories, industry multiples'],
        ['src/constants/initialData.ts', '115', 'Default sample company data and assumptions'],
        ['src/constants/valuationStyles.ts', '~35', 'Conservative/Moderate/Aggressive style presets'],
        ['src/hooks/useFinancialData.ts', '215', 'State management with undo/redo and auto-save'],
        ['src/hooks/useValuationCalculations.ts', '209', 'Memoized calculation pipeline with scenario/style adjustments'],
        ['src/hooks/useHistory.ts', '~80', '50-state undo/redo history stack'],
        ['src/hooks/useTheme.ts', '~40', 'Dark/light mode with localStorage persistence'],
        ['src/hooks/useKeyboardShortcuts.ts', '~45', 'Ctrl+Z/Y/S keyboard bindings'],
        ['src/services/stockAPI.ts', '729', 'FMP Stable API client with peer company support'],
        ['src/services/yahooFinanceAPI.ts', '~350', 'Yahoo Finance fallback client'],
        ['src/workers/monteCarlo.worker.ts', '143', 'Web Worker for 5,000 Monte Carlo simulations'],
        ['src/App.tsx', '177', 'Main orchestrator with lazy-loaded tabs'],
    ]
)

# ── Footer ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("— End of Document —")
r.italic = True
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("WOLF Valuation Engine v2.0 — Complete Technical Overview — March 2026")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
output_path = os.path.join(os.path.dirname(__file__), 'WOLF_Engine_Complete_Overview_v2.docx')
doc.save(output_path)
print(f"✅ Document saved: {output_path}")
print(f"   Sections: {len(toc_items)}")
